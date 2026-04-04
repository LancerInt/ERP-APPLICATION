from io import BytesIO
from decimal import Decimal
from rest_framework import viewsets, status, permissions
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.exceptions import ValidationError
from django.http import HttpResponse
from rbac.permissions import HasModulePermission
from django.db.models import Q, Sum
from django.shortcuts import get_object_or_404

from .models import (
    CustomerPOUpload,
    SalesOrder,
    DispatchChallan,
    SalesInvoiceCheck,
    SalesFreightDetail,
    FreightAdviceOutbound,
    FreightPayment,
    ReceivableLedger,
)
from .serializers import (
    CustomerPOUploadSerializer,
    CreateCustomerPOSerializer,
    SalesOrderListSerializer,
    SalesOrderDetailSerializer,
    CreateSalesOrderSerializer,
    DispatchChallanListSerializer,
    DispatchChallanDetailSerializer,
    CreateUpdateDCSerializer,
    SalesInvoiceListSerializer,
    SalesInvoiceDetailSerializer,
    CreateUpdateSalesInvoiceSerializer,
    SalesFreightDetailListSerializer,
    SalesFreightDetailDetailSerializer,
    CreateUpdateFreightDetailSerializer,
    FreightAdviceOutboundListSerializer,
    FreightAdviceOutboundDetailSerializer,
    CreateUpdateFreightSerializer,
    FreightPaymentSerializer,
    FreightPaymentListSerializer,
    FreightAttachmentSerializer,
    ReceivableLedgerListSerializer,
    ReceivableLedgerDetailSerializer,
    CreateUpdateReceivableSerializer,
)
from .services import (
    POUploadService,
    SalesOrderService,
    DispatchService,
    InvoiceService,
    FreightService,
    ReceivableService,
)
from .selectors import (
    SalesOrderSelector,
    DispatchSelector,
    ReceivableSelector,
    SalesReconciliationSelector,
)


class CustomerPOUploadViewSet(viewsets.ModelViewSet):
    """
    ViewSet for Customer PO.
    Full CRUD with nested line items.
    """

    queryset = CustomerPOUpload.objects.all()
    permission_classes = [permissions.IsAuthenticated, HasModulePermission]
    module_name = 'Customer PO'
    filterset_fields = ['customer', 'company', 'status']
    search_fields = ['upload_id', 'po_number', 'customer__customer_name']
    ordering_fields = ['upload_date', 'po_date', 'status']
    ordering = ['-upload_date']

    def get_queryset(self):
        return (
            CustomerPOUpload.objects
            .select_related('customer', 'company', 'warehouse', 'linked_sales_order')
            .prefetch_related('parsed_lines__parsed_sku')
            .filter(is_active=True)
        )

    def get_serializer_class(self):
        if self.action in ('create', 'update', 'partial_update'):
            return CreateCustomerPOSerializer
        return CustomerPOUploadSerializer

    def perform_destroy(self, instance):
        """Delete Customer PO and rollback: unlink from any SOs."""
        from django.db import transaction as db_transaction
        with db_transaction.atomic():
            # Unlink from M2M SOs
            for so in instance.linked_sales_orders.all():
                so.customer_pos.remove(instance)
            # Soft delete
            instance.is_active = False
            instance.save(update_fields=['is_active', 'updated_at'])

    @action(detail=True, methods=['post'])
    def trigger_parsing(self, request, pk=None):
        """Trigger AI parsing of uploaded PO"""
        po_upload = self.get_object()

        if po_upload.status != 'UPLOADED':
            raise ValidationError(f"Cannot parse PO in {po_upload.status} status")

        POUploadService.upload_and_parse_customer_po(po_upload)

        return Response(
            {'detail': 'Parsing queued successfully'},
            status=status.HTTP_202_ACCEPTED
        )

    @action(detail=True, methods=['post'])
    def convert_to_sales_order(self, request, pk=None):
        """Convert parsed PO to sales order"""
        po_upload = self.get_object()

        if po_upload.status != 'PARSED':
            raise ValidationError("PO must be in PARSED status to convert")

        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)

        try:
            sales_order = SalesOrderService.create_sales_order_from_parsed_po(
                po_upload=po_upload,
                company_id=serializer.validated_data['company'],
                warehouse_id=serializer.validated_data['warehouse'],
                price_list_id=serializer.validated_data['price_list'],
                credit_terms=serializer.validated_data.get('credit_terms', ''),
                freight_terms=serializer.validated_data.get('freight_terms', ''),
                required_ship_date=serializer.validated_data.get('required_ship_date'),
            )

            return Response(
                {
                    'detail': 'Sales order created',
                    'sales_order_id': sales_order.id,
                    'sales_order_no': sales_order.so_no,
                },
                status=status.HTTP_201_CREATED
            )
        except Exception as e:
            raise ValidationError(f"Failed to create sales order: {str(e)}")

    @action(detail=True, methods=['post'])
    def request_manual_review(self, request, pk=None):
        """Request manual review for problematic PO"""
        po_upload = self.get_object()
        review_comments = request.data.get('review_comments', '')

        po_upload.manual_review_required = True
        po_upload.review_comments = review_comments
        po_upload.save()

        return Response(
            {'detail': 'Manual review requested'},
            status=status.HTTP_200_OK
        )


class SalesOrderViewSet(viewsets.ModelViewSet):
    """
    ViewSet for sales orders.
    Manages SO creation, approval, and line item tracking.
    """

    queryset = SalesOrder.objects.all()
    permission_classes = [permissions.IsAuthenticated, HasModulePermission]
    module_name = 'Sales Order'
    filterset_fields = ['customer', 'warehouse', 'approval_status', 'company']
    search_fields = ['so_no', 'customer__name']
    ordering_fields = ['so_date', 'approval_status']
    ordering = ['-so_date']

    def get_queryset(self):
        return (
            SalesOrder.objects
            .select_related('customer', 'company', 'warehouse', 'price_list', 'approved_by')
            .prefetch_related('so_lines__product')
            .filter(is_active=True)
        )

    def get_serializer_class(self):
        if self.action in ('create', 'update', 'partial_update'):
            return CreateSalesOrderSerializer
        elif self.action == 'retrieve':
            return SalesOrderDetailSerializer
        return SalesOrderListSerializer

    def perform_destroy(self, instance):
        """Delete SO and cascade the full chain:
        SO → DCs → Freight Details → Outward Freights → Payments, Invoices → Receivables.
        Also rollback reserved_qty and reset Customer POs."""
        from django.db import transaction as db_transaction
        from .models import DCLine

        with db_transaction.atomic():
            # 1. Find all DCs linked to this SO
            dc_ids = DCLine.objects.filter(
                linked_so_line__so=instance
            ).values_list('dc_id', flat=True).distinct()

            for dc in DispatchChallan.objects.filter(id__in=dc_ids, is_active=True):
                # Cascade: Receivables ← Invoices ← DC
                for inv in dc.invoice_checks.filter(is_active=True):
                    for recv in inv.receivables.filter(is_active=True):
                        recv.is_active = False
                        recv.save(update_fields=['is_active', 'updated_at'])
                    inv.is_active = False
                    inv.save(update_fields=['is_active', 'updated_at'])

                # Cascade: Freight Details linked to this DC
                fd_ids = dc.freight_detail_links.values_list('freight_detail_id', flat=True).distinct()
                for fd in SalesFreightDetail.objects.filter(id__in=fd_ids, is_active=True):
                    for of in fd.outward_freights.filter(is_active=True):
                        of.payments.all().delete()
                        of.is_active = False
                        of.save(update_fields=['is_active', 'updated_at'])
                    fd.is_active = False
                    fd.save(update_fields=['is_active', 'updated_at'])
                dc.freight_detail_links.all().delete()

                # Cascade: Outward Freights directly linked to this DC
                for of in dc.freight_advices.filter(is_active=True):
                    of.payments.all().delete()
                    of.is_active = False
                    of.save(update_fields=['is_active', 'updated_at'])

                # Reverse reserved_qty for each DC line
                for dc_line in dc.dc_lines.select_related('linked_so_line').all():
                    if dc_line.linked_so_line:
                        so_line = dc_line.linked_so_line
                        so_line.reserved_qty = max(Decimal('0'), so_line.reserved_qty - dc_line.quantity_dispatched)
                        so_line.save(update_fields=['reserved_qty', 'updated_at'])

                dc.is_active = False
                dc.save(update_fields=['is_active', 'updated_at'])

            # 2. Unlink Customer POs (M2M) - set them back to CONFIRMED
            for po in instance.customer_pos.all():
                if po.status == 'CONVERTED':
                    po.status = 'CONFIRMED'
                    po.save(update_fields=['status', 'updated_at'])
            instance.customer_pos.clear()

            # 3. Soft delete the SO
            instance.is_active = False
            instance.save(update_fields=['is_active', 'updated_at'])

    @action(detail=False, methods=['get'])
    def next_so_number(self, request):
        """Get the next auto-generated SO number"""
        so_no = SalesOrderService._generate_so_number()
        return Response({'so_no': so_no})

    @action(detail=True, methods=['post'])
    def approve(self, request, pk=None):
        """Approve sales order and auto-create Dispatch Challan."""
        from django.db import transaction as db_transaction
        from .models import DispatchChallan, DCLine

        sales_order = self.get_object()

        try:
            approved_by = request.user.stakeholder_profile
        except (AttributeError, Exception):
            approved_by = None

        with db_transaction.atomic():
            SalesOrderService.approve_sales_order(sales_order, approved_by)

            # Auto-create DC from SO lines
            so_lines = sales_order.so_lines.select_related('product').all()
            dc_no = ''
            dc_id = ''
            if so_lines.exists():
                dc_no = DispatchService._generate_dc_number()
                dc = DispatchChallan.objects.create(
                    dc_no=dc_no,
                    warehouse=sales_order.warehouse,
                    status='DRAFT',
                    created_by=request.user,
                )
                for so_line in so_lines:
                    pending = so_line.get_pending_qty()
                    if pending > 0:
                        DCLine.objects.create(
                            dc=dc,
                            product=so_line.product,
                            quantity_dispatched=pending,
                            uom=so_line.uom,
                            linked_so_line=so_line,
                        )
                dc_id = str(dc.id)

        return Response(
            {
                'detail': f'Sales order approved. DC {dc_no} created.',
                'so_no': sales_order.so_no,
                'dc_no': dc_no,
                'dc_id': dc_id,
            },
            status=status.HTTP_200_OK
        )

    @action(detail=True, methods=['post'])
    def reject(self, request, pk=None):
        """Reject sales order"""
        sales_order = self.get_object()

        SalesOrderService.reject_sales_order(sales_order)

        return Response(
            {'detail': 'Sales order rejected'},
            status=status.HTTP_200_OK
        )

    @action(detail=True, methods=['get'], url_path='download-pdf')
    def download_pdf(self, request, pk=None):
        """Generate Sales Order PDF — professional Tally-style landscape voucher"""
        so = self.get_object()
        lines = so.so_lines.select_related('product').order_by('line_no')

        def fmt(val):
            n = float(val or 0)
            return f'{n:,.2f}'

        # --- Address helper ---
        def addr_str(addr):
            if not addr:
                return ''
            if isinstance(addr, list):
                return ', '.join(str(a) for a in addr if a)
            if isinstance(addr, dict):
                parts = [addr.get('street', ''), addr.get('city', ''), addr.get('state', ''), addr.get('postal_code', '')]
                return ', '.join(p for p in parts if p)
            return str(addr)

        # --- Amount in words (Indian) ---
        def amount_in_words(amount):
            try:
                amt_int = int(round(float(amount)))
                if amt_int == 0:
                    return 'Zero'
                ones = ['', 'One', 'Two', 'Three', 'Four', 'Five', 'Six', 'Seven', 'Eight', 'Nine',
                        'Ten', 'Eleven', 'Twelve', 'Thirteen', 'Fourteen', 'Fifteen', 'Sixteen',
                        'Seventeen', 'Eighteen', 'Nineteen']
                tens = ['', '', 'Twenty', 'Thirty', 'Forty', 'Fifty', 'Sixty', 'Seventy', 'Eighty', 'Ninety']
                def two_d(n):
                    if n < 20: return ones[n]
                    return tens[n // 10] + (' ' + ones[n % 10] if n % 10 else '')
                def three_d(n):
                    if n >= 100:
                        return ones[n // 100] + ' Hundred' + (' and ' + two_d(n % 100) if n % 100 else '')
                    return two_d(n)
                if amt_int >= 10000000:
                    return three_d(amt_int // 10000000) + ' Crore ' + amount_in_words(amt_int % 10000000)
                if amt_int >= 100000:
                    return two_d(amt_int // 100000) + ' Lakh ' + amount_in_words(amt_int % 100000)
                if amt_int >= 1000:
                    return two_d(amt_int // 1000) + ' Thousand ' + amount_in_words(amt_int % 1000)
                return three_d(amt_int)
            except Exception:
                return ''

        category_map = {
            'RAW_MATERIAL': 'Raw Material', 'PACKING_MATERIAL': 'Packing Material',
            'FINISHED_GOOD': 'Finished Good', 'SEMI_FINISHED': 'Semi Finished',
            'TRADED_PRODUCTS': 'Traded Products', 'CAPITAL_GOOD': 'Capital Good',
            'MACHINE_SPARES': 'Machine Spares', 'CONSUMABLES': 'Consumables',
        }
        term_map = {
            'TO_PAY': 'To Pay', 'PAID': 'Paid', 'NET_15': 'Net 15',
            'NET_30': 'Net 30', 'NET_45': 'Net 45', 'CUSTOM': 'Custom',
        }

        # --- Compute ---
        line_rows = ''
        total_qty = Decimal('0')
        subtotal = Decimal('0')
        total_discount = Decimal('0')
        total_taxable = Decimal('0')
        total_tax = Decimal('0')
        grand_total = Decimal('0')
        total_pending = Decimal('0')

        for i, line in enumerate(lines, 1):
            qty = line.quantity_ordered or Decimal('0')
            price = line.unit_price or Decimal('0')
            disc = line.discount or Decimal('0')
            gst = line.gst or Decimal('0')
            gross = qty * price
            taxable = gross - disc
            tax = taxable * gst / Decimal('100')
            net = taxable + tax
            pending = line.get_pending_qty()
            cat = category_map.get(line.product.goods_sub_type, '-')
            del_date = line.delivery_schedule_date.strftime('%d-%b-%Y') if line.delivery_schedule_date else ''

            total_qty += qty
            subtotal += gross
            total_discount += disc
            total_taxable += taxable
            total_tax += tax
            grand_total += net
            total_pending += pending

            line_rows += f'''<tr>
<td class="c">{i}</td>
<td class="l">{cat}</td>
<td class="l"><b>{line.product.product_name}</b><br/><span class="sku">{line.product.sku_code}</span></td>
<td class="r">{fmt(qty)}</td>
<td class="c">{line.uom}</td>
<td class="r">{fmt(price)}</td>
<td class="r">{fmt(gross)}</td>
<td class="r">{fmt(disc)}</td>
<td class="r">{fmt(taxable)}</td>
<td class="c">{gst}%</td>
<td class="r">{fmt(tax)}</td>
<td class="r b">{fmt(net)}</td>
<td class="r pend">{fmt(pending)}</td>
<td class="c">{del_date}</td>
</tr>'''

        company = so.company
        customer = so.customer
        co_name = company.legal_name if company else '-'
        co_gstin = company.gstin or ''
        co_addr = addr_str(company.registered_address) if company else ''
        co_phone = company.contact_phone or ''
        co_email = company.contact_email or ''

        cu_name = customer.customer_name if customer else '-'
        cu_code = customer.customer_code if customer else ''
        cu_gstin = customer.gstin or ''
        cu_addr = addr_str(customer.billing_address) if customer else ''
        cu_phone = customer.contact_phone or ''
        cu_contact = customer.contact_person if customer else ''
        cu_email = customer.contact_email if customer else ''

        so_date = so.so_date.strftime('%d-%b-%Y') if so.so_date else '-'
        ship_date = so.required_ship_date.strftime('%d-%b-%Y') if so.required_ship_date else '-'
        wh_name = so.warehouse.name if so.warehouse else '-'
        pl_name = so.price_list.price_list_id if so.price_list else '-'

        words = amount_in_words(grand_total).strip()
        paise = int(round((float(grand_total) - int(float(grand_total))) * 100))
        amt_words = f'Rupees {words}'
        if paise > 0:
            amt_words += f' and {amount_in_words(paise).strip()} Paise'
        amt_words += ' Only'

        html = f'''<!DOCTYPE html>
<html>
<head><meta charset="utf-8"/>
<style>
@page {{ size: A4 landscape; margin: 10mm; }}
body {{ font-family: "Segoe UI", Arial, sans-serif; font-size: 9px; color: #1a1a1a; margin: 0; }}
.page {{ border: 2px solid #222; }}

/* === HEADER === */
.hdr {{ background: #1a3a5c; color: #fff; padding: 12px 20px; }}
.hdr h1 {{ margin: 0; font-size: 18px; letter-spacing: 1.5px; text-transform: uppercase; }}
.hdr-sub {{ font-size: 9px; color: #cbd5e1; margin-top: 3px; }}
.hdr-right {{ float: right; text-align: right; font-size: 9px; color: #e2e8f0; padding-top: 2px; }}

/* === TITLE BAR === */
.title-bar {{ background: #f0f4f8; border-bottom: 2px solid #222; padding: 6px 20px;
              font-size: 13px; font-weight: 700; text-align: center; text-transform: uppercase;
              letter-spacing: 2px; color: #1a3a5c; }}

/* === INFO SECTION === */
.info {{ border-bottom: 1px solid #ccc; padding: 0; }}
.info table {{ width: 100%; border-collapse: collapse; }}
.info td {{ padding: 4px 12px; font-size: 9px; vertical-align: top; }}
.info .lbl {{ font-weight: 700; color: #475569; width: 90px; white-space: nowrap; }}
.info .val {{ color: #111; }}
.info .val b {{ font-size: 10px; }}
.info .div {{ border-right: 1px solid #ccc; }}
.info .section-label {{ font-size: 8px; text-transform: uppercase; letter-spacing: 1px;
                        color: #94a3b8; font-weight: 700; padding: 6px 12px 2px; }}

/* === TERMS BAR === */
.terms-bar {{ background: #f8fafc; border-bottom: 1px solid #ccc; padding: 5px 20px; }}
.terms-bar table {{ width: 100%; }}
.terms-bar td {{ font-size: 9px; padding: 2px 0; }}
.terms-bar .lbl {{ font-weight: 700; color: #475569; }}

/* === ITEMS TABLE === */
.items {{ width: 100%; border-collapse: collapse; }}
.items th {{ background: #1a3a5c; color: #fff; padding: 5px 4px; font-size: 7.5px;
             text-transform: uppercase; letter-spacing: 0.5px; font-weight: 700;
             border: 1px solid #0f2a44; white-space: nowrap; }}
.items td {{ padding: 5px 4px; font-size: 8.5px; border-bottom: 1px solid #e2e8f0;
             border-left: 1px solid #e2e8f0; border-right: 1px solid #e2e8f0; }}
.items tr:nth-child(even) td {{ background: #f8fafc; }}
.items td.c {{ text-align: center; }}
.items td.r {{ text-align: right; }}
.items td.l {{ text-align: left; }}
.items td.b {{ font-weight: 700; }}
.items td.pend {{ color: #c2410c; font-weight: 600; }}
.items .sku {{ font-size: 7.5px; color: #64748b; }}

/* Footer row */
.items tfoot td {{ background: #f0f4f8; border: 1px solid #222; font-weight: 700;
                   font-size: 9px; padding: 5px 4px; }}

/* === SUMMARY === */
.summary {{ border-top: 2px solid #222; }}
.summary table {{ width: 100%; border-collapse: collapse; }}
.summary td {{ padding: 3px 20px; font-size: 9px; }}
.summary .s-lbl {{ text-align: right; width: 82%; color: #475569; }}
.summary .s-val {{ text-align: right; width: 18%; font-weight: 600; border-left: 1px solid #ccc; }}
.summary .gt td {{ font-size: 11px; font-weight: 700; color: #1a3a5c;
                   border-top: 2px solid #222; padding: 5px 20px; }}

/* === WORDS === */
.words {{ border-top: 1px solid #ccc; padding: 5px 20px; font-size: 9px; }}
.words em {{ font-weight: 700; font-style: italic; }}

/* === REMARKS === */
.remarks {{ border-top: 1px solid #ccc; padding: 5px 20px; font-size: 8.5px; color: #475569; }}

/* === SIGNATURE === */
.sig {{ border-top: 2px solid #222; }}
.sig table {{ width: 100%; border-collapse: collapse; }}
.sig td {{ padding: 10px 25px; font-size: 9px; vertical-align: bottom; }}
.sig .line {{ border-top: 1px solid #333; margin-top: 35px; padding-top: 5px; text-align: center;
              font-size: 8.5px; color: #475569; }}
</style>
</head>
<body>
<div class="page">

<!-- HEADER -->
<div class="hdr">
    <span class="hdr-right">
        {f'GSTIN: {co_gstin}' if co_gstin else ''}<br/>
        {f'Phone: {co_phone}' if co_phone else ''}
        {f' | Email: {co_email}' if co_email else ''}
    </span>
    <h1>{co_name}</h1>
    <div class="hdr-sub">{co_addr}</div>
</div>

<!-- TITLE -->
<div class="title-bar">Sales Order</div>

<!-- INFO -->
<div class="info">
    <table>
        <tr>
            <td colspan="2" class="section-label">Order Information</td>
            <td colspan="2" class="section-label">Buyer Details</td>
        </tr>
        <tr>
            <td class="lbl div" style="width:12%;">Voucher No.</td>
            <td class="val div" style="width:25%;"><b>{so.so_no}</b></td>
            <td class="lbl" style="width:12%;">Customer</td>
            <td class="val" style="width:25%;"><b>{cu_name}</b> ({cu_code})</td>
        </tr>
        <tr>
            <td class="lbl div">Date</td>
            <td class="val div"><b>{so_date}</b></td>
            <td class="lbl">GSTIN</td>
            <td class="val">{cu_gstin or '-'}</td>
        </tr>
        <tr>
            <td class="lbl div">Warehouse</td>
            <td class="val div">{wh_name}</td>
            <td class="lbl">Address</td>
            <td class="val">{cu_addr or '-'}</td>
        </tr>
        <tr>
            <td class="lbl div">Price List</td>
            <td class="val div">{pl_name}</td>
            <td class="lbl">Contact</td>
            <td class="val">{cu_contact or '-'}{f' | Ph: {cu_phone}' if cu_phone else ''}{f' | {cu_email}' if cu_email else ''}</td>
        </tr>
        <tr>
            <td class="lbl div">Status</td>
            <td class="val div"><b>{so.approval_status}</b></td>
            <td class="lbl">Ship Date</td>
            <td class="val"><b>{ship_date}</b></td>
        </tr>
    </table>
</div>

<!-- TERMS -->
<div class="terms-bar">
    <table><tr>
        <td><span class="lbl">Freight Terms:</span> {term_map.get(so.freight_terms, so.freight_terms or '-')}</td>
        <td><span class="lbl">Credit Terms:</span> {term_map.get(so.credit_terms, so.credit_terms or '-')}</td>
        <td><span class="lbl">Delivery Date:</span> {ship_date}</td>
        <td><span class="lbl">PO Ref:</span> {so.customer_po_reference or '-'}</td>
    </tr></table>
</div>

<!-- ITEMS -->
<table class="items">
<thead><tr>
    <th style="width:18px;">Sl</th>
    <th style="width:70px;">Category</th>
    <th>Particulars</th>
    <th style="width:48px;text-align:right;">Qty</th>
    <th style="width:28px;">UOM</th>
    <th style="width:52px;text-align:right;">Rate</th>
    <th style="width:56px;text-align:right;">Gross Amt</th>
    <th style="width:42px;text-align:right;">Disc.</th>
    <th style="width:56px;text-align:right;">Taxable</th>
    <th style="width:28px;">GST</th>
    <th style="width:48px;text-align:right;">Tax Amt</th>
    <th style="width:60px;text-align:right;">Net Amt</th>
    <th style="width:48px;text-align:right;color:#fdba74;">Pending</th>
    <th style="width:52px;">Del. Date</th>
</tr></thead>
<tbody>{line_rows}</tbody>
<tfoot><tr>
    <td colspan="3" class="r">Total</td>
    <td class="r">{fmt(total_qty)}</td>
    <td></td>
    <td></td>
    <td class="r">{fmt(subtotal)}</td>
    <td class="r">{fmt(total_discount)}</td>
    <td class="r">{fmt(total_taxable)}</td>
    <td></td>
    <td class="r">{fmt(total_tax)}</td>
    <td class="r" style="font-size:10px;">{fmt(grand_total)}</td>
    <td class="r pend">{fmt(total_pending)}</td>
    <td></td>
</tr></tfoot>
</table>

<!-- SUMMARY -->
<div class="summary">
    <table>
        <tr><td class="s-lbl">Sub Total (Gross)</td><td class="s-val">{fmt(subtotal)}</td></tr>
        <tr><td class="s-lbl">(-) Total Discount</td><td class="s-val">{fmt(total_discount)}</td></tr>
        <tr><td class="s-lbl">Taxable Amount</td><td class="s-val">{fmt(total_taxable)}</td></tr>
        <tr><td class="s-lbl">GST Amount</td><td class="s-val">{fmt(total_tax)}</td></tr>
        <tr class="gt"><td class="s-lbl">Grand Total</td><td class="s-val">{fmt(grand_total)}</td></tr>
    </table>
</div>

<!-- AMOUNT IN WORDS -->
<div class="words"><em>Amount in Words:</em> {amt_words}</div>

<!-- REMARKS -->
{f'<div class="remarks"><b>Remarks:</b> {so.remarks}</div>' if so.remarks else ''}

<!-- SIGNATURE -->
<div class="sig">
    <table><tr>
        <td style="width:33%;"><div class="line">Prepared By</div></td>
        <td style="width:33%;"><div class="line">Receiver\'s Signature</div></td>
        <td style="width:34%;text-align:right;"><div class="line">Authorised Signatory<br/><b>{co_name}</b></div></td>
    </tr></table>
</div>

</div>
</body></html>'''

        try:
            from xhtml2pdf import pisa
            buffer = BytesIO()
            pisa.CreatePDF(html, dest=buffer)
            buffer.seek(0)
            response = HttpResponse(buffer.read(), content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{so.so_no}.pdf"'
            return response
        except ImportError:
            return Response({'error': 'PDF generation library not available'}, status=500)

    @action(detail=True, methods=['get'], url_path='dispatch-lines')
    def dispatch_lines(self, request, pk=None):
        """Get SO lines with qty info for DC creation."""
        so = self.get_object()
        lines = so.so_lines.select_related('product').all()
        result = []
        for line in lines:
            result.append({
                'so_line_id': str(line.id),
                'line_no': line.line_no,
                'product': str(line.product.id),
                'product_name': line.product.product_name,
                'product_sku': line.product.sku_code,
                'product_category': line.product.goods_sub_type or '',
                'quantity_ordered': str(line.quantity_ordered),
                'reserved_qty': str(line.reserved_qty),
                'pending_qty': str(line.get_pending_qty()),
                'uom': line.uom,
            })
        return Response({
            'so_id': str(so.id),
            'so_no': so.so_no,
            'customer_name': so.customer.customer_name if so.customer else '',
            'company_name': so.company.legal_name if so.company else '',
            'warehouse': str(so.warehouse.id) if so.warehouse else '',
            'warehouse_name': so.warehouse.name if so.warehouse else '',
            'destination': so.destination or '',
            'freight_terms': so.freight_terms or '',
            'lines': result,
        })

    @action(detail=True, methods=['get'], url_path='linked-dcs')
    def linked_dcs(self, request, pk=None):
        """Get Dispatch Challans linked to this Sales Order"""
        so = self.get_object()
        from .models import DispatchChallan, DCLine
        dc_ids = DCLine.objects.filter(
            linked_so_line__so=so
        ).values_list('dc_id', flat=True).distinct()
        dcs = DispatchChallan.objects.filter(id__in=dc_ids).values(
            'id', 'dc_no', 'dispatch_date', 'status'
        )
        return Response(list(dcs))

    @action(detail=True, methods=['get'], url_path='validate-flow')
    def validate_flow(self, request, pk=None):
        """Flow Bot: Validate the entire sales chain for this SO and report discrepancies."""
        so = self.get_object()
        from .models import (
            DCLine, DispatchChallan, SalesFreightDetail, FreightDetailDCLink,
            FreightAdviceOutbound, SalesInvoiceCheck, SalesInvoiceLine,
        )
        issues = []
        warnings = []
        info = []

        so_lines = list(so.so_lines.select_related('product').all())

        # ═══════════ 1. CPO → SO Validation ═══════════
        cpos = list(so.customer_pos.prefetch_related('parsed_lines').all())
        if not cpos:
            warnings.append({'step': 'CPO → SO', 'field': 'Customer PO', 'message': 'No Customer PO linked to this Sales Order'})
        else:
            for cpo in cpos:
                cpo_lines = list(cpo.parsed_lines.all())
                for cl in cpo_lines:
                    matching_so = [sl for sl in so_lines if sl.product_id == cl.parsed_sku_id]
                    if not matching_so:
                        issues.append({'step': 'CPO → SO', 'field': 'Product',
                            'message': f'CPO {cpo.upload_id} has product "{cl.product_description}" not found in SO lines'})
                        continue
                    sl = matching_so[0]
                    cpo_qty = cl.quantity or Decimal('0')
                    so_qty = sl.quantity_ordered or Decimal('0')
                    if abs(cpo_qty - so_qty) > Decimal('0.01'):
                        issues.append({'step': 'CPO → SO', 'field': 'Quantity',
                            'message': f'Product "{cl.product_description}": CPO qty={cpo_qty} ≠ SO qty={so_qty}'})
                    cpo_rate = cl.price or Decimal('0')
                    so_rate = sl.unit_price or Decimal('0')
                    if abs(cpo_rate - so_rate) > Decimal('0.01'):
                        issues.append({'step': 'CPO → SO', 'field': 'Price',
                            'message': f'Product "{cl.product_description}": CPO rate=₹{cpo_rate} ≠ SO rate=₹{so_rate}'})
                    if cl.uom and sl.uom and cl.uom != sl.uom:
                        issues.append({'step': 'CPO → SO', 'field': 'UOM',
                            'message': f'Product "{cl.product_description}": CPO UOM={cl.uom} ≠ SO UOM={sl.uom}'})

        # ═══════════ 2. SO → DC Validation ═══════════
        dc_ids = DCLine.objects.filter(
            linked_so_line__so=so
        ).values_list('dc_id', flat=True).distinct()
        dcs = list(DispatchChallan.objects.filter(id__in=dc_ids, is_active=True).prefetch_related('dc_lines__product', 'dc_lines__linked_so_line'))

        if not dcs:
            warnings.append({'step': 'SO → DC', 'field': 'Dispatch Challan', 'message': 'No Dispatch Challans linked to this SO'})
        else:
            info.append({'step': 'SO → DC', 'field': 'Count', 'message': f'{len(dcs)} Dispatch Challan(s) found'})
            for sl in so_lines:
                total_dispatched = Decimal('0')
                for dc in dcs:
                    for dcl in dc.dc_lines.all():
                        if dcl.linked_so_line_id == sl.id:
                            total_dispatched += dcl.quantity_dispatched or Decimal('0')
                            if dcl.product_id != sl.product_id:
                                issues.append({'step': 'SO → DC', 'field': 'Product',
                                    'message': f'DC {dc.dc_no}: DC line product ≠ SO line product for line #{sl.line_no}'})
                            if dcl.uom and sl.uom and dcl.uom != sl.uom:
                                issues.append({'step': 'SO → DC', 'field': 'UOM',
                                    'message': f'DC {dc.dc_no}: UOM mismatch for {sl.product.product_name if sl.product else "?"} — DC={dcl.uom}, SO={sl.uom}'})
                so_qty = sl.quantity_ordered or Decimal('0')
                if total_dispatched > so_qty + Decimal('0.01'):
                    issues.append({'step': 'SO → DC', 'field': 'Quantity',
                        'message': f'{sl.product.product_name if sl.product else "?"}: Dispatched={total_dispatched} exceeds SO qty={so_qty}'})
                elif total_dispatched < so_qty - Decimal('0.01'):
                    warnings.append({'step': 'SO → DC', 'field': 'Quantity',
                        'message': f'{sl.product.product_name if sl.product else "?"}: Dispatched={total_dispatched} < SO qty={so_qty} (pending dispatch)'})

        # ═══════════ 3. DC → Freight Details Validation ═══════════
        fd_links = FreightDetailDCLink.objects.filter(dc__in=dc_ids).select_related('freight_detail', 'dc')
        fd_ids = set(fl.freight_detail_id for fl in fd_links)
        freight_details = list(SalesFreightDetail.objects.filter(id__in=fd_ids, is_active=True))

        if dcs and not freight_details:
            warnings.append({'step': 'DC → Freight', 'field': 'Freight Details', 'message': 'No Freight Details created for DCs'})
        elif freight_details:
            info.append({'step': 'DC → Freight', 'field': 'Count', 'message': f'{len(freight_details)} Freight Detail(s) found'})
            for fd in freight_details:
                if fd.total_freight and fd.total_freight <= 0:
                    warnings.append({'step': 'DC → Freight', 'field': 'Amount', 'message': f'Freight {fd.freight_no}: Total freight is ₹0'})
                if fd.balance_freight and fd.balance_freight < 0:
                    issues.append({'step': 'DC → Freight', 'field': 'Balance', 'message': f'Freight {fd.freight_no}: Negative balance ₹{fd.balance_freight}'})

        # ═══════════ 4. Freight → Outward Freight Validation ═══════════
        outward_freights = list(FreightAdviceOutbound.objects.filter(
            freight_detail__in=fd_ids, is_active=True
        ).prefetch_related('payments'))

        if freight_details and not outward_freights:
            warnings.append({'step': 'Freight → Outward', 'field': 'Outward Freight', 'message': 'No Outward Freight records created'})
        elif outward_freights:
            info.append({'step': 'Freight → Outward', 'field': 'Count', 'message': f'{len(outward_freights)} Outward Freight(s) found'})
            for of in outward_freights:
                payable = of.payable_amount or Decimal('0')
                paid = of.total_paid or Decimal('0')
                balance = of.balance or Decimal('0')
                if abs(payable - paid - balance) > Decimal('0.01'):
                    issues.append({'step': 'Freight → Outward', 'field': 'Payment',
                        'message': f'{of.advice_no}: Payable(₹{payable}) - Paid(₹{paid}) ≠ Balance(₹{balance})'})
                # Cross-check with Freight Detail
                if of.freight_detail_id and of.freight_detail_id in fd_ids:
                    fd = next((f for f in freight_details if f.id == of.freight_detail_id), None)
                    if fd and fd.total_freight and payable and abs(fd.total_freight - payable) > Decimal('1'):
                        warnings.append({'step': 'Freight → Outward', 'field': 'Amount',
                            'message': f'{of.advice_no}: Outward payable ₹{payable} ≠ Freight Detail total ₹{fd.total_freight}'})

        # ═══════════ 5. Invoice Validation ═══════════
        invoices = list(SalesInvoiceCheck.objects.filter(
            Q(dc_reference__in=dc_ids) | Q(so_reference=so), is_active=True
        ).prefetch_related('invoice_lines'))

        if dcs and not invoices:
            warnings.append({'step': 'Invoice', 'field': 'Invoice', 'message': 'No Invoices created for this flow'})
        elif invoices:
            info.append({'step': 'Invoice', 'field': 'Count', 'message': f'{len(invoices)} Invoice(s) found'})
            for inv in invoices:
                lines = list(inv.invoice_lines.all())
                calc_subtotal = sum(l.amount for l in lines)
                if abs(calc_subtotal - (inv.subtotal or Decimal('0'))) > Decimal('0.01'):
                    issues.append({'step': 'Invoice', 'field': 'Subtotal',
                        'message': f'{inv.invoice_no}: Line items total ₹{calc_subtotal} ≠ Invoice subtotal ₹{inv.subtotal}'})
                calc_cgst = sum(l.cgst_amount for l in lines)
                calc_sgst = sum(l.sgst_amount for l in lines)
                calc_igst = sum(l.igst_amount for l in lines)
                if abs(calc_cgst - (inv.cgst_total or Decimal('0'))) > Decimal('0.5'):
                    issues.append({'step': 'Invoice', 'field': 'CGST',
                        'message': f'{inv.invoice_no}: Calculated CGST ₹{calc_cgst} ≠ Stored ₹{inv.cgst_total}'})
                if abs(calc_sgst - (inv.sgst_total or Decimal('0'))) > Decimal('0.5'):
                    issues.append({'step': 'Invoice', 'field': 'SGST',
                        'message': f'{inv.invoice_no}: Calculated SGST ₹{calc_sgst} ≠ Stored ₹{inv.sgst_total}'})
                # Check products exist in SO
                for il in lines:
                    if il.product_id:
                        so_product_ids = [sl.product_id for sl in so_lines]
                        if il.product_id not in so_product_ids:
                            warnings.append({'step': 'Invoice', 'field': 'Product',
                                'message': f'{inv.invoice_no}: Invoice product "{il.description}" not found in SO lines'})

        # ═══════════ 6. Receivable Validation ═══════════
        from .models import ReceivableLedger
        inv_ids = [inv.id for inv in invoices]
        receivables = list(ReceivableLedger.objects.filter(invoice_reference__in=inv_ids, is_active=True))

        if invoices and not receivables:
            warnings.append({'step': 'Receivable', 'field': 'Receivable', 'message': 'No Receivables created for invoices'})
        elif receivables:
            info.append({'step': 'Receivable', 'field': 'Count', 'message': f'{len(receivables)} Receivable(s) found'})
            for r in receivables:
                inv = next((i for i in invoices if i.id == r.invoice_reference_id), None)
                if inv and abs((r.amount or Decimal('0')) - (inv.grand_total or Decimal('0'))) > Decimal('0.01'):
                    issues.append({'step': 'Receivable', 'field': 'Amount',
                        'message': f'Receivable for {inv.invoice_no}: Amount ₹{r.amount} ≠ Invoice total ₹{inv.grand_total}'})
                balance = (r.amount or Decimal('0')) - (r.amount_paid or Decimal('0'))
                if abs(balance - (r.balance or Decimal('0'))) > Decimal('0.01'):
                    issues.append({'step': 'Receivable', 'field': 'Balance',
                        'message': f'Receivable: Calculated balance ₹{balance} ≠ Stored balance ₹{r.balance}'})

        # ═══════════ Summary ═══════════
        total_checks = len(issues) + len(warnings) + len(info)
        return Response({
            'so_no': so.so_no,
            'customer_name': so.customer.customer_name if so.customer else '',
            'total_checks': total_checks,
            'issues_count': len(issues),
            'warnings_count': len(warnings),
            'info_count': len(info),
            'status': 'PASS' if len(issues) == 0 else 'FAIL',
            'issues': issues,
            'warnings': warnings,
            'info': info,
        })

    @action(detail=True, methods=['post'], url_path='flow-chat')
    def flow_chat(self, request, pk=None):
        """Smart Flow Bot chat — gathers full flow data and answers any question."""
        so = self.get_object()
        question = request.data.get('message', '').strip()
        if not question:
            return Response({'answer': 'Please ask a question about this sales flow.'})

        from .models import (
            DCLine, DispatchChallan, SalesFreightDetail, FreightDetailDCLink,
            FreightAdviceOutbound, SalesInvoiceCheck, SalesInvoiceLine, ReceivableLedger,
        )

        # ── Gather ALL flow data ──
        so_lines = list(so.so_lines.select_related('product').all())
        cpos = list(so.customer_pos.prefetch_related('parsed_lines').all())
        dc_ids = list(DCLine.objects.filter(linked_so_line__so=so).values_list('dc_id', flat=True).distinct())
        dcs = list(DispatchChallan.objects.filter(id__in=dc_ids, is_active=True).prefetch_related('dc_lines__product'))
        fd_links = FreightDetailDCLink.objects.filter(dc__in=dc_ids).select_related('freight_detail')
        fd_ids = set(fl.freight_detail_id for fl in fd_links)
        freight_details = list(SalesFreightDetail.objects.filter(id__in=fd_ids, is_active=True))
        outward_freights = list(FreightAdviceOutbound.objects.filter(
            freight_detail__in=fd_ids, is_active=True
        ).prefetch_related('payments'))
        invoices = list(SalesInvoiceCheck.objects.filter(
            Q(dc_reference__in=dc_ids) | Q(so_reference=so), is_active=True
        ).prefetch_related('invoice_lines'))
        inv_ids = [inv.id for inv in invoices]
        receivables = list(ReceivableLedger.objects.filter(invoice_reference__in=inv_ids, is_active=True)) if inv_ids else []

        # ── Build DETAILED data context ──
        ctx = {
            'so': {
                'so_no': so.so_no, 'status': so.approval_status, 'date': str(so.so_date),
                'customer': so.customer.customer_name if so.customer else '',
                'company': so.company.legal_name if so.company else '',
                'warehouse': so.warehouse.name if so.warehouse else '',
                'destination': so.destination or '',
                'freight_terms': so.freight_terms or '', 'payment_terms': so.payment_terms or '',
                'currency': so.currency or 'INR',
                'delivery_terms': getattr(so, 'delivery_terms', '') or '',
                'delivery_location': getattr(so, 'delivery_location', '') or '',
                'delivery_due_date': str(so.delivery_due_date) if getattr(so, 'delivery_due_date', None) else '',
                'required_ship_date': str(so.required_ship_date) if so.required_ship_date else '',
                'party_code': getattr(so, 'party_code', '') or '',
                'indent_no': getattr(so, 'indent_no', '') or '',
                'indent_date': str(so.indent_date) if getattr(so, 'indent_date', None) else '',
                'dispatched_through': getattr(so, 'dispatched_through', '') or '',
                'consignee_name': getattr(so, 'consignee_name', '') or '',
                'consignee_address': getattr(so, 'consignee_address', '') or '',
                'consignee_gstin': getattr(so, 'consignee_gstin', '') or '',
                'billing_address': getattr(so, 'billing_address', '') or '',
                'billing_gstin': getattr(so, 'billing_gstin', '') or '',
                'special_instructions': getattr(so, 'special_instructions', '') or '',
                'remarks': so.remarks or '',
                'total_amount': str(so.get_total_amount()),
                'total_lines': len(so_lines),
                'lines': [{
                    'line_no': sl.line_no, 'product': sl.product.product_name if sl.product else '',
                    'sku': sl.product.sku_code if sl.product else '',
                    'category': sl.product.goods_sub_type if sl.product else '',
                    'hsn': sl.product.hsn_code if sl.product else '',
                    'qty_ordered': str(sl.quantity_ordered), 'uom': sl.uom, 'unit_price': str(sl.unit_price),
                    'discount': str(sl.discount), 'gst_percent': str(sl.gst),
                    'line_total': str(sl.get_line_total()),
                    'qty_reserved': str(sl.reserved_qty), 'qty_pending': str(sl.get_pending_qty()),
                    'delivery_date': str(sl.delivery_schedule_date) if sl.delivery_schedule_date else '',
                } for sl in so_lines],
            },
            'cpos': [{
                'po_no': cpo.upload_id, 'customer_po_number': cpo.po_number or '',
                'status': cpo.status, 'date': str(cpo.po_date) if cpo.po_date else '',
                'company': cpo.company.legal_name if cpo.company else '',
                'customer': cpo.customer.customer_name if cpo.customer else '',
                'warehouse': cpo.warehouse.name if cpo.warehouse else '',
                'delivery_type': cpo.delivery_type or '', 'payment_terms': cpo.payment_terms or '',
                'freight_terms': cpo.freight_terms or '', 'currency': cpo.currency or 'INR',
                'party_code': cpo.party_code or '', 'indent_no': cpo.indent_no or '',
                'consignee_name': cpo.consignee_name or '', 'consignee_address': cpo.consignee_address or '',
                'consignee_gstin': cpo.consignee_gstin or '',
                'billing_address': cpo.billing_address or '', 'billing_gstin': cpo.billing_gstin or '',
                'destination': cpo.destination or '', 'delivery_location': cpo.delivery_location or '',
                'special_instructions': cpo.special_instructions or '', 'remarks': cpo.remarks or '',
                'total_lines': cpo.parsed_lines.count(),
                'lines': [{
                    'product': l.product_description or '',
                    'sku': l.parsed_sku.sku_code if l.parsed_sku else '',
                    'item_code': l.item_code or '', 'hsn_code': l.hsn_code or '',
                    'qty': str(l.quantity), 'uom': l.uom, 'price': str(l.price),
                    'discount': str(l.discount), 'gst': str(l.gst),
                    'sgst': str(l.sgst_percent), 'cgst': str(l.cgst_percent), 'igst': str(l.igst_percent),
                    'line_total': str(l.line_total) if hasattr(l, 'line_total') and l.line_total else '',
                    'delivery_date': str(l.delivery_schedule_date) if l.delivery_schedule_date else '',
                } for l in cpo.parsed_lines.all()],
            } for cpo in cpos],
            'dcs': [{
                'dc_no': dc.dc_no, 'status': dc.status,
                'dispatch_date': str(dc.dispatch_date) if dc.dispatch_date else '',
                'warehouse': dc.warehouse.name if dc.warehouse else '',
                'invoice_no': dc.invoice_no or '', 'invoice_date': str(dc.invoice_date) if dc.invoice_date else '',
                'lorry_no': dc.lorry_no or '', 'transporter': dc.transporter.name if dc.transporter else '',
                'destination': getattr(dc, 'destination', '') or '',
                'total_dispatch_qty': str(sum(dl.quantity_dispatched or 0 for dl in dc.dc_lines.all())),
                'total_lines': dc.dc_lines.count(),
                'lines': [{
                    'product': dl.product.product_name if dl.product else '',
                    'sku': dl.product.sku_code if dl.product else '',
                    'qty_dispatched': str(dl.quantity_dispatched), 'uom': dl.uom,
                    'batch': dl.batch or '', 'noa': str(dl.noa) if dl.noa else '',
                    'weight': str(dl.weight) if dl.weight else '',
                    'so_line_no': str(dl.linked_so_line.line_no) if dl.linked_so_line else '',
                    'so_qty_ordered': str(dl.linked_so_line.quantity_ordered) if dl.linked_so_line else '',
                } for dl in dc.dc_lines.all()],
            } for dc in dcs],
            'freight_details': [{
                'freight_no': fd.freight_no, 'status': fd.status,
                'date': str(fd.freight_date) if fd.freight_date else '',
                'company': fd.company.legal_name if fd.company else '',
                'customer': fd.customer.customer_name if fd.customer else '',
                'factory': fd.factory.name if fd.factory else '',
                'transporter': fd.transporter.name if fd.transporter else '',
                'freight_type': fd.freight_type or '', 'lorry_no': fd.lorry_no or '',
                'total_quantity': str(fd.total_quantity) if fd.total_quantity else '0',
                'quantity_uom': fd.quantity_uom or '',
                'freight_per_ton': str(fd.freight_per_ton) if fd.freight_per_ton else '0',
                'total_freight': str(fd.total_freight), 'paid': str(fd.freight_paid),
                'balance': str(fd.balance_freight),
                'destination': fd.destination or '', 'destination_state': fd.destination_state or '',
                'discount': str(fd.discount) if hasattr(fd, 'discount') and fd.discount else '0',
                'additional_freight': str(fd.additional_freight) if hasattr(fd, 'additional_freight') and fd.additional_freight else '0',
                'dc_links': [{
                    'dc_no': fl.dc.dc_no if fl.dc else '', 'qty': str(fl.quantity),
                    'invoice_no': fl.invoice_no or '', 'destination': fl.destination or '',
                } for fl in fd.dc_links.all()],
            } for fd in freight_details],
            'outward_freights': [{
                'advice_no': of.advice_no, 'status': of.status,
                'freight_date': str(of.freight_date) if of.freight_date else '',
                'invoice_date': str(of.invoice_date) if of.invoice_date else '',
                'customer': of.customer_name or '', 'lorry_no': of.lorry_no or '',
                'destination': of.destination or '',
                'shipment_quantity': str(of.shipment_quantity) if of.shipment_quantity else '0',
                'quantity_uom': of.quantity_uom or '',
                'base_amount': str(of.base_amount), 'freight_per_ton': str(of.freight_per_ton),
                'payable_amount': str(of.payable_amount), 'total_paid': str(of.total_paid),
                'balance': str(of.balance),
                'remarks': of.remarks or '',
                'payment_count': of.payments.count(),
                'payments': [{
                    'date': str(p.payment_date), 'amount': str(p.amount_paid),
                    'mode': p.get_payment_mode_display() if hasattr(p, 'get_payment_mode_display') else p.payment_mode,
                    'reference': p.reference_no or '', 'remarks': p.remarks or '',
                } for p in of.payments.all()],
                'dc_links': [{
                    'dc_no': dl.dc.dc_no if dl.dc else '', 'invoice_no': dl.invoice_no or '',
                    'destination': dl.destination or '',
                } for dl in of.dc_links.all()],
            } for of in outward_freights],
            'invoices': [{
                'invoice_no': inv.invoice_no, 'status': inv.status,
                'date': str(inv.invoice_date) if inv.invoice_date else '',
                'company': inv.company.legal_name if inv.company else '',
                'company_gstin': inv.company_gstin or '',
                'customer': inv.customer.customer_name if inv.customer else '',
                'consignee_name': inv.consignee_name or '', 'consignee_address': inv.consignee_address or '',
                'consignee_gstin': inv.consignee_gstin or '',
                'buyer_name': inv.buyer_name or '', 'buyer_address': inv.buyer_address or '',
                'buyer_gstin': inv.buyer_gstin or '',
                'dc_no': inv.dc_reference.dc_no if inv.dc_reference else '',
                'so_no': inv.so_reference.so_no if inv.so_reference else '',
                'buyers_order_no': inv.buyers_order_no or '',
                'dispatch_doc_no': inv.dispatch_doc_no or '',
                'destination': inv.destination or '', 'terms_of_delivery': inv.terms_of_delivery or '',
                'payment_terms': inv.payment_terms or '',
                'subtotal': str(inv.subtotal), 'cgst_total': str(inv.cgst_total),
                'sgst_total': str(inv.sgst_total), 'igst_total': str(inv.igst_total),
                'round_off': str(inv.round_off), 'grand_total': str(inv.grand_total),
                'amount_in_words': inv.amount_in_words or '',
                'declaration': inv.declaration or '', 'remarks': inv.remarks or '',
                'total_lines': inv.invoice_lines.count(),
                'lines': [{
                    'sl_no': il.sl_no, 'description': il.description,
                    'hsn_sac': il.hsn_sac or '', 'qty': str(il.quantity), 'uom': il.uom or '',
                    'rate': str(il.rate), 'discount_percent': str(il.discount_percent),
                    'amount': str(il.amount), 'gst_rate': str(il.gst_rate),
                    'cgst_rate': str(il.cgst_rate), 'cgst_amount': str(il.cgst_amount),
                    'sgst_rate': str(il.sgst_rate), 'sgst_amount': str(il.sgst_amount),
                    'igst_rate': str(il.igst_rate), 'igst_amount': str(il.igst_amount),
                } for il in inv.invoice_lines.all()],
            } for inv in invoices],
            'receivables': [{
                'invoice_no': r.invoice_reference.invoice_no if r.invoice_reference else '',
                'customer': r.customer.customer_name if r.customer else '',
                'payment_status': r.payment_status,
                'invoice_date': str(r.invoice_date) if r.invoice_date else '',
                'due_date': str(r.due_date) if r.due_date else '',
                'total_amount': str(r.amount), 'amount_received': str(r.amount_paid),
                'balance_due': str(r.balance),
                'is_overdue': r.due_date and r.due_date < __import__('datetime').date.today() and r.payment_status != 'PAID',
                'notes': r.notes or '',
            } for r in receivables],
            'flow_summary': {
                'total_cpos': len(cpos), 'total_so_lines': len(so_lines),
                'total_dcs': len(dcs), 'total_freight_details': len(freight_details),
                'total_outward_freights': len(outward_freights), 'total_invoices': len(invoices),
                'total_receivables': len(receivables),
                'so_total_value': str(so.get_total_amount()),
                'total_dispatched_qty': str(sum(
                    sum(dl.quantity_dispatched or 0 for dl in dc.dc_lines.all()) for dc in dcs
                )),
                'total_freight_value': str(sum(float(fd.total_freight or 0) for fd in freight_details)),
                'total_freight_paid': str(sum(float(fd.freight_paid or 0) for fd in freight_details)),
                'total_invoice_value': str(sum(float(inv.grand_total or 0) for inv in invoices)),
                'total_receivable_amount': str(sum(float(r.amount or 0) for r in receivables)),
                'total_received': str(sum(float(r.amount_paid or 0) for r in receivables)),
                'total_outstanding': str(sum(float(r.balance or 0) for r in receivables)),
            },
        }

        # ── Intelligent Query Engine ──
        import json, re
        q = question.lower()
        data_str = json.dumps(ctx, indent=2, default=str)

        # Try Claude API first
        try:
            import anthropic
            client = anthropic.Anthropic()
            system_prompt = (
                "You are a Sales Flow Validation Bot for an Indian ERP system (Lancer ERP). "
                "You have access to the complete sales flow data for a specific Sales Order. "
                "Answer the user's question accurately using ONLY the data provided. "
                "Use ₹ for currency, Indian number format. Be concise but thorough. "
                "If data is missing or empty, say so. Use bullet points and bold for clarity. "
                "The flow chain is: Customer PO → Sales Order → Dispatch Challan → Freight Details → Outward Freight → Invoice → Receivable.\n\n"
                f"FLOW DATA:\n{data_str}"
            )
            response = client.messages.create(
                model="claude-haiku-4-5-20251001",
                max_tokens=1024,
                system=system_prompt,
                messages=[{"role": "user", "content": question}],
            )
            answer = response.content[0].text
            return Response({'answer': answer, 'source': 'ai'})
        except Exception:
            pass

        # ── Fallback: Smart data-driven response engine ──
        answer = self._smart_answer(q, ctx)
        return Response({'answer': answer, 'source': 'engine'})

    @staticmethod
    def _smart_answer(q, ctx):
        """Data-driven answer engine when AI is not available."""
        so = ctx['so']
        cpos = ctx['cpos']
        dcs = ctx['dcs']
        fds = ctx['freight_details']
        ofs = ctx['outward_freights']
        invs = ctx['invoices']
        recs = ctx['receivables']
        lines = so['lines']

        # ── Pattern matchers ──
        import re

        # Totals / amounts
        if re.search(r'total|amount|value|worth|grand|how much', q):
            r = f"📊 **Financial Summary for {so['so_no']}**\n\n"
            r += f"• SO Total: ₹{so['total_amount']}\n"
            if invs:
                for inv in invs:
                    r += f"• Invoice {inv['invoice_no']}: Subtotal ₹{inv['subtotal']} | CGST ₹{inv['cgst']} | SGST ₹{inv['sgst']} | Grand Total ₹{inv['grand_total']}\n"
            if recs:
                total_recv = sum(float(r['paid']) for r in recs)
                total_bal = sum(float(r['balance']) for r in recs)
                r += f"• Total Received: ₹{total_recv:,.2f} | Outstanding: ₹{total_bal:,.2f}\n"
            if ofs:
                total_freight = sum(float(o['payable']) for o in ofs)
                total_fpaid = sum(float(o['paid']) for o in ofs)
                r += f"• Freight Payable: ₹{total_freight:,.2f} | Freight Paid: ₹{total_fpaid:,.2f}\n"
            return r

        # Products / items / what products
        if re.search(r'product|item|what.*sell|what.*order|goods|material|sku', q):
            r = f"📦 **Products in {so['so_no']}** ({len(lines)} items)\n\n"
            for i, l in enumerate(lines, 1):
                r += f"{i}. **{l['product']}** ({l['sku']})\n   Qty: {l['qty']} {l['uom']} | Rate: ₹{l['price']} | Total: ₹{l['total']}\n   Dispatched/Reserved: {l['reserved']} | Pending: {l['pending']}\n\n"
            return r

        # Quantity / dispatch status / pending
        if re.search(r'quantit|qty|dispatch|pending|remain|how many|shipped|deliver', q):
            r = f"📊 **Dispatch Status for {so['so_no']}**\n\n"
            for l in lines:
                status = '✅ Fully dispatched' if float(l['pending']) <= 0 else f"⏳ Pending: {l['pending']} {l['uom']}"
                r += f"• **{l['product']}**: Ordered {l['qty']} {l['uom']} | Reserved {l['reserved']} | {status}\n"
            r += f"\n**DCs:** {len(dcs)} dispatch challan(s)\n"
            for dc in dcs:
                r += f"  • {dc['dc_no']} [{dc['status']}] — {len(dc['lines'])} item(s)\n"
            return r

        # Customer / who
        if re.search(r'customer|who|buyer|client|consignee', q):
            r = f"👤 **Customer Details**\n\n"
            r += f"• Customer: **{so['customer']}**\n"
            r += f"• Company: {so['company']}\n"
            r += f"• Warehouse: {so['warehouse']}\n"
            if invs and invs[0].get('buyer'):
                r += f"• Buyer (Invoice): {invs[0]['buyer']}\n"
                r += f"• Consignee: {invs[0]['consignee']}\n"
            return r

        # Status / where is it / stage
        if re.search(r'status|stage|where.*is|progress|state|lifecycle|flow', q):
            r = f"📋 **Flow Status for {so['so_no']}**\n\n"
            r += f"• SO Status: **{so['status']}**\n"
            r += f"• Customer POs: {len(cpos)} linked — {', '.join(c['po_no'] for c in cpos) or 'None'}\n"
            dc_str = ', '.join(d['dc_no'] + ' [' + d['status'] + ']' for d in dcs) or 'None'
            r += f"• Dispatch Challans: {len(dcs)} — {dc_str}\n"
            fd_str = ', '.join(f['freight_no'] + ' [' + f['status'] + ']' for f in fds) or 'None'
            r += f"• Freight Details: {len(fds)} — {fd_str}\n"
            of_str = ', '.join(o['advice_no'] + ' [' + o['status'] + ']' for o in ofs) or 'None'
            r += f"• Outward Freight: {len(ofs)} — {of_str}\n"
            inv_str = ', '.join(i['invoice_no'] + ' [' + i['status'] + ']' for i in invs) or 'None'
            r += f"• Invoices: {len(invs)} — {inv_str}\n"
            rec_str = ', '.join(rc['invoice_no'] + ' [' + rc['status'] + ']' for rc in recs) or 'None'
            r += f"• Receivables: {len(recs)} — {rec_str}\n"
            # Next action suggestion
            if not dcs: r += "\n💡 **Next:** Create Dispatch Challans"
            elif not fds: r += "\n💡 **Next:** Create Freight Details"
            elif not invs: r += "\n💡 **Next:** Create Invoice"
            elif not recs: r += "\n💡 **Next:** Create Receivable"
            else:
                unpaid = [r for r in recs if r['status'] != 'PAID']
                if unpaid: r += "\n💡 **Next:** Collect pending receivables"
                else: r += "\n✅ **Flow complete!**"
            return r

        # Price / rate
        if re.search(r'price|rate|cost|how much.*per|unit price', q):
            r = f"💰 **Pricing for {so['so_no']}**\n\n"
            for l in lines:
                r += f"• **{l['product']}**: ₹{l['price']} per {l['uom']} | Disc: ₹{l['discount']} | GST: {l['gst']}%\n"
            if cpos:
                r += f"\n**CPO Rates (for comparison):**\n"
                for cpo in cpos:
                    for cl in cpo['lines']:
                        r += f"• {cl['product']}: ₹{cl['price']} per {cl['uom']}\n"
            return r

        # Invoice / GST / tax
        if re.search(r'invoice|gst|tax|cgst|sgst|igst|bill', q):
            if not invs:
                return f"🧾 No invoices have been created for {so['so_no']} yet."
            r = f"🧾 **Invoice Details**\n\n"
            for inv in invs:
                r += f"**{inv['invoice_no']}** [{inv['status']}] — {inv['date']}\n"
                r += f"  Subtotal: ₹{inv['subtotal']} | CGST: ₹{inv['cgst']} | SGST: ₹{inv['sgst']} | IGST: ₹{inv['igst']}\n"
                r += f"  **Grand Total: ₹{inv['grand_total']}**\n"
                if inv['lines']:
                    r += f"  Items: {len(inv['lines'])}\n"
                    for il in inv['lines']:
                        r += f"    • {il['description']} — Qty:{il['qty']} × ₹{il['rate']} = ₹{il['amount']} (GST {il['gst_rate']}%)\n"
                r += '\n'
            return r

        # Freight / transport / lorry
        if re.search(r'freight|transport|lorry|truck|shipping cost', q):
            if not fds and not ofs:
                return f"🚛 No freight records found for {so['so_no']}."
            r = f"🚛 **Freight Details**\n\n"
            for fd in fds:
                r += f"• **{fd['freight_no']}** [{fd['status']}]\n  Total: ₹{fd['total_freight']} | Paid: ₹{fd['paid']} | Balance: ₹{fd['balance']}\n  Lorry: {fd['lorry_no'] or '-'} | Dest: {fd['destination'] or '-'} | Per Ton: ₹{fd['freight_per_ton']}\n\n"
            if ofs:
                r += "**Outward Freight:**\n"
                for o in ofs:
                    r += f"• **{o['advice_no']}** [{o['status']}] — Payable: ₹{o['payable']} | Paid: ₹{o['paid']} | Balance: ₹{o['balance']}\n"
                    if o['payments']:
                        for p in o['payments']:
                            r += f"  💳 {p['date']}: ₹{p['amount']} ({p['mode']})\n"
            return r

        # Payment / receivable / collection / outstanding / due
        if re.search(r'pay|receiv|collect|outstand|due|balance|owe', q):
            if not recs:
                return f"💳 No receivable records found for {so['so_no']}. Create an invoice and receivable first."
            r = f"💳 **Receivables for {so['so_no']}**\n\n"
            total_amt = total_paid = total_bal = 0
            for rec in recs:
                amt, paid, bal = float(rec['amount']), float(rec['paid']), float(rec['balance'])
                total_amt += amt; total_paid += paid; total_bal += bal
                status_icon = '✅' if rec['status'] == 'PAID' else '⏳' if rec['status'] == 'PARTIALLY_PAID' else '🔴'
                r += f"{status_icon} **{rec['invoice_no']}** [{rec['status']}]\n   Amount: ₹{amt:,.2f} | Received: ₹{paid:,.2f} | Balance: ₹{bal:,.2f}"
                if rec['due_date']: r += f" | Due: {rec['due_date']}"
                r += '\n\n'
            r += f"**Summary:** Total ₹{total_amt:,.2f} | Received ₹{total_paid:,.2f} | Outstanding ₹{total_bal:,.2f}"
            return r

        # CPO / customer PO
        if re.search(r'cpo|customer po|purchase order', q):
            if not cpos:
                return f"📄 No Customer POs linked to {so['so_no']}."
            r = f"📄 **Customer POs linked to {so['so_no']}**\n\n"
            for cpo in cpos:
                r += f"• **{cpo['po_no']}** (PO# {cpo['po_number'] or '-'}) [{cpo['status']}] — {cpo['date']}\n"
                for cl in cpo['lines']:
                    r += f"  → {cl['product']}: {cl['qty']} {cl['uom']} × ₹{cl['price']}\n"
                r += '\n'
            return r

        # DC / dispatch challan
        if re.search(r'\bdc\b|dispatch|challan', q):
            if not dcs:
                return f"🚚 No Dispatch Challans for {so['so_no']}."
            r = f"🚚 **Dispatch Challans for {so['so_no']}**\n\n"
            for dc in dcs:
                r += f"• **{dc['dc_no']}** [{dc['status']}] — {dc['date']}\n  Invoice: {dc['invoice_no'] or '-'} | Lorry: {dc['lorry_no'] or '-'}\n"
                for dl in dc['lines']:
                    r += f"  → {dl['product']}: {dl['qty_dispatched']} {dl['uom']}\n"
                r += '\n'
            return r

        # UOM
        if re.search(r'uom|unit|measure|mts|kilogram', q):
            r = f"📏 **UOM Details for {so['so_no']}**\n\n"
            for l in lines:
                r += f"• {l['product']}: **{l['uom']}** (Qty: {l['qty']})\n"
            return r

        # Comparison / mismatch / difference / verify
        if re.search(r'compar|mismatch|differ|verify|check|match|valid', q):
            return "💡 Use the **⚡ Re-run** button to perform a full validation check. It will compare quantities, prices, UOMs, and amounts across the entire chain."

        # Summary / overview / tell me about / everything
        if re.search(r'summar|overview|tell me|everything|all|detail|full|report|info', q):
            r = f"📋 **Complete Flow Summary — {so['so_no']}**\n\n"
            r += f"**Customer:** {so['customer']} | **Company:** {so['company']}\n"
            r += f"**Status:** {so['status']} | **Date:** {so['date']}\n"
            r += f"**SO Value:** ₹{so['total_amount']} | **Items:** {len(lines)}\n\n"
            r += f"**Chain:**\n"
            r += f"  📄 {len(cpos)} Customer PO(s)\n"
            r += f"  📦 {len(lines)} SO Line(s)\n"
            r += f"  🚚 {len(dcs)} Dispatch Challan(s)\n"
            r += f"  📦 {len(fds)} Freight Detail(s)\n"
            r += f"  🚛 {len(ofs)} Outward Freight(s)\n"
            r += f"  🧾 {len(invs)} Invoice(s)\n"
            r += f"  💰 {len(recs)} Receivable(s)\n"
            return r

        # Greeting
        if re.search(r'^(hi|hello|hey|good|thanks|thank)', q):
            return f"👋 Hello! I'm your Flow Bot for **{so['so_no']}**.\n\nI have full access to this order's data — {len(lines)} products, {len(dcs)} DCs, {len(invs)} invoices, and more.\n\nAsk me anything! For example:\n• \"Show me the products\"\n• \"What's the dispatch status?\"\n• \"How much is outstanding?\"\n• \"Any freight payments made?\""

        # Help
        if re.search(r'help|what can|how to|command', q):
            return ("💡 **I can answer questions like:**\n\n"
                "• \"What products are in this order?\"\n"
                "• \"How much is the total?\"\n"
                "• \"What's the dispatch status?\"\n"
                "• \"Show me the invoices\"\n"
                "• \"Any payments received?\"\n"
                "• \"Who is the customer?\"\n"
                "• \"What's the freight cost?\"\n"
                "• \"Show me the CPO details\"\n"
                "• \"What stage is this order at?\"\n"
                "• \"Price breakdown\"\n"
                "• \"Summary of everything\"\n\n"
                "Just ask naturally — I understand the full sales flow!")

        # Fallback — search all data for keywords
        import json
        data_str = json.dumps(ctx, default=str).lower()
        words = [w for w in q.split() if len(w) > 2]
        matches = []
        for word in words:
            if word in data_str:
                matches.append(word)
        if matches:
            return (f"I found references to **{', '.join(matches)}** in the flow data. "
                f"Could you be more specific? For example:\n"
                f"• \"Show me the {matches[0]} details\"\n"
                f"• \"What's the status of {matches[0]}?\"\n\n"
                f"Or type **help** to see what I can answer.")

        return (f"I have full access to the sales flow for **{so['so_no']}** "
            f"({len(lines)} products, {len(dcs)} DCs, {len(invs)} invoices).\n\n"
            f"Try asking:\n"
            f"• \"Show me the products\"\n"
            f"• \"What's the total amount?\"\n"
            f"• \"Dispatch status?\"\n"
            f"• \"Invoice details\"\n"
            f"• \"Any outstanding payments?\"")

    @action(detail=False, methods=['get'])
    def pending_for_warehouse(self, request):
        """Get pending SOs for specified warehouse"""
        warehouse_id = request.query_params.get('warehouse_id')

        if not warehouse_id:
            raise ValidationError("warehouse_id query parameter required")

        sos = SalesOrderSelector.get_pending_sos_for_warehouse(warehouse_id)
        serializer = SalesOrderListSerializer(sos, many=True)

        return Response(serializer.data)

    @action(detail=False, methods=['get'])
    def approved_not_dispatched(self, request):
        """Get approved SOs with pending dispatch"""
        warehouse_id = request.query_params.get('warehouse_id')

        sos = SalesOrderSelector.get_approved_sos_not_dispatched(warehouse_id)
        serializer = SalesOrderListSerializer(sos, many=True)

        return Response(serializer.data)


class DispatchChallanViewSet(viewsets.ModelViewSet):
    """
    ViewSet for dispatch challans.
    Manages DC creation, release, and delivery tracking.
    """

    queryset = DispatchChallan.objects.all()
    permission_classes = [permissions.IsAuthenticated, HasModulePermission]
    module_name = 'Dispatch Challan'
    filterset_fields = ['warehouse', 'transporter', 'status']
    search_fields = ['dc_no']
    ordering_fields = ['dispatch_date', 'status']
    ordering = ['-dispatch_date']

    def get_queryset(self):
        return (
            DispatchChallan.objects
            .select_related('warehouse', 'transporter', 'freight_advice_link')
            .prefetch_related(
                'dc_lines__product',
                'dc_lines__linked_so_line__so__customer',
                'dc_lines__linked_so_line__so__company',
                'delivery_locations',
            )
            .filter(is_active=True)
        )

    def get_serializer_class(self):
        if self.action in ('create', 'update', 'partial_update'):
            return CreateUpdateDCSerializer
        elif self.action == 'retrieve':
            return DispatchChallanDetailSerializer
        return DispatchChallanListSerializer

    def perform_destroy(self, instance):
        """Delete DC and cascade: reverse SO qty, delete downstream Freight Details,
        Outward Freights, Payments, Invoices, and Receivables."""
        from django.db import transaction as db_transaction

        with db_transaction.atomic():
            # 1. Cascade delete Receivables linked via Invoices referencing this DC
            for inv in instance.invoice_checks.filter(is_active=True):
                for recv in inv.receivables.filter(is_active=True):
                    recv.is_active = False
                    recv.save(update_fields=['is_active', 'updated_at'])
                inv.is_active = False
                inv.save(update_fields=['is_active', 'updated_at'])

            # 2. Cascade delete Freight Details that link to this DC
            fd_ids = instance.freight_detail_links.values_list('freight_detail_id', flat=True).distinct()
            for fd in SalesFreightDetail.objects.filter(id__in=fd_ids, is_active=True):
                # Delete Outward Freights linked to this Freight Detail
                for of in fd.outward_freights.filter(is_active=True):
                    of.payments.all().delete()
                    of.is_active = False
                    of.save(update_fields=['is_active', 'updated_at'])
                fd.is_active = False
                fd.save(update_fields=['is_active', 'updated_at'])
            # Remove the DC link records
            instance.freight_detail_links.all().delete()

            # 3. Cascade delete Outward Freights directly linked to this DC
            for of in instance.freight_advices.filter(is_active=True):
                of.payments.all().delete()
                if of.freight_detail and of.freight_detail.is_active:
                    self._recalc_freight_detail(of.freight_detail, exclude_of=of)
                of.is_active = False
                of.save(update_fields=['is_active', 'updated_at'])

            # 4. Reverse reserved_qty for all DC lines linked to SO lines
            affected_sos = set()
            for dc_line in instance.dc_lines.select_related('linked_so_line__so').all():
                if dc_line.linked_so_line:
                    so_line = dc_line.linked_so_line
                    so_line.reserved_qty = max(
                        Decimal('0'), so_line.reserved_qty - dc_line.quantity_dispatched
                    )
                    so_line.save(update_fields=['reserved_qty', 'updated_at'])
                    affected_sos.add(so_line.so)

            # 5. Update SO dispatch status
            for so in affected_sos:
                so.update_dispatch_status()

            # 6. Soft delete the DC
            instance.is_active = False
            instance.save(update_fields=['is_active', 'updated_at'])

    @staticmethod
    def _recalc_freight_detail(fd, exclude_of=None):
        """Recalculate FreightDetail paid/balance after outward freight changes."""
        qs = FreightAdviceOutbound.objects.filter(freight_detail=fd, is_active=True)
        if exclude_of:
            qs = qs.exclude(id=exclude_of.id)
        total = qs.aggregate(total=Sum('total_paid'))['total'] or Decimal('0')
        fd.freight_paid = total
        fd.balance_freight = max(Decimal('0'), (fd.total_freight or Decimal('0')) - total)
        if fd.balance_freight <= 0 and total > 0:
            fd.status = 'COMPLETED'
        elif total > 0:
            fd.status = 'IN_PROGRESS'
        else:
            fd.status = 'PENDING'
        fd.save(update_fields=['freight_paid', 'balance_freight', 'status', 'updated_at'])

    @action(detail=True, methods=['post'])
    def release(self, request, pk=None):
        """Release DC for dispatch"""
        dc = self.get_object()

        DispatchService.release_dispatch_challan(dc)

        return Response(
            {'detail': 'Dispatch challan released', 'dc_no': dc.dc_no},
            status=status.HTTP_200_OK
        )

    @action(detail=True, methods=['post'])
    def mark_delivered(self, request, pk=None):
        """Mark DC as delivered"""
        dc = self.get_object()

        DispatchService.mark_dispatch_delivered(dc)

        return Response(
            {'detail': 'Dispatch challan marked delivered'},
            status=status.HTTP_200_OK
        )

    @action(detail=False, methods=['get'])
    def open_dcs(self, request):
        """Get open (not delivered) dispatch challans"""
        warehouse_id = request.query_params.get('warehouse_id')

        dcs = DispatchSelector.get_open_dcs(warehouse_id)
        serializer = DispatchChallanListSerializer(dcs, many=True)

        return Response(serializer.data)

    @action(detail=False, methods=['get'])
    def awaiting_invoice(self, request):
        """Get delivered DCs awaiting invoice"""
        warehouse_id = request.query_params.get('warehouse_id')

        dcs = DispatchSelector.get_delivered_dcs_without_invoice(warehouse_id)
        serializer = DispatchChallanListSerializer(dcs, many=True)

        return Response(serializer.data)


class SalesInvoiceCheckViewSet(viewsets.ModelViewSet):
    """ViewSet for GST Sales Invoices."""
    queryset = SalesInvoiceCheck.objects.all()
    permission_classes = [permissions.IsAuthenticated, HasModulePermission]
    module_name = 'Sales Invoice'
    search_fields = ['invoice_no', 'customer__customer_name', 'buyer_name']
    ordering_fields = ['invoice_date', 'grand_total']
    ordering = ['-invoice_date']

    def get_queryset(self):
        return (
            SalesInvoiceCheck.objects
            .select_related('company', 'customer', 'dc_reference', 'so_reference')
            .prefetch_related('invoice_lines')
            .filter(is_active=True)
        )

    def get_serializer_class(self):
        if self.action in ('create', 'update', 'partial_update'):
            return CreateUpdateSalesInvoiceSerializer
        elif self.action == 'retrieve':
            return SalesInvoiceDetailSerializer
        return SalesInvoiceListSerializer

    def perform_destroy(self, instance):
        """Soft delete invoice and cascade delete linked receivables."""
        from django.db import transaction as db_transaction
        with db_transaction.atomic():
            for recv in instance.receivables.filter(is_active=True):
                recv.is_active = False
                recv.save(update_fields=['is_active', 'updated_at'])
            instance.is_active = False
            instance.save(update_fields=['is_active', 'updated_at'])


class SalesFreightDetailViewSet(viewsets.ModelViewSet):
    """
    Freight Details ViewSet — parent entry for outward freight.
    """
    queryset = SalesFreightDetail.objects.all()
    permission_classes = [permissions.IsAuthenticated, HasModulePermission]
    module_name = 'Freight Advice'
    filterset_fields = ['company', 'factory', 'status', 'customer', 'transporter']
    search_fields = ['freight_no', 'lorry_no', 'destination']
    ordering_fields = ['freight_date', 'status', 'total_freight']
    ordering = ['-freight_date']

    def get_queryset(self):
        return (
            SalesFreightDetail.objects
            .select_related('company', 'factory', 'customer', 'transporter')
            .prefetch_related('dc_links__dc')
            .filter(is_active=True)
        )

    def get_serializer_class(self):
        if self.action in ('create', 'update', 'partial_update'):
            return CreateUpdateFreightDetailSerializer
        elif self.action == 'retrieve':
            return SalesFreightDetailDetailSerializer
        return SalesFreightDetailListSerializer

    def perform_destroy(self, instance):
        """Delete Freight Detail and cascade: delete linked Outward Freights and their Payments."""
        from django.db import transaction as db_transaction
        with db_transaction.atomic():
            # Cascade delete Outward Freights and their payments
            for of in instance.outward_freights.filter(is_active=True):
                of.payments.all().delete()
                of.is_active = False
                of.save(update_fields=['is_active', 'updated_at'])
            # Soft delete
            instance.is_active = False
            instance.save(update_fields=['is_active', 'updated_at'])

    @action(detail=False, methods=['get'], url_path='available-dcs')
    def available_dcs(self, request):
        """Get DCs not yet linked to any Freight Detail.
        Optional: ?customer_id=xxx to filter by customer (via SO link)."""
        from .models import FreightDetailDCLink, DCLine
        linked_dc_ids = FreightDetailDCLink.objects.values_list('dc_id', flat=True).distinct()
        qs = DispatchChallan.objects.filter(is_active=True).exclude(id__in=linked_dc_ids)

        # Filter by customer if provided (customer comes from linked SO)
        customer_id = request.query_params.get('customer_id')
        if customer_id:
            # Find SO line IDs for this customer's SOs
            from .models import SalesOrder
            customer_so_ids = SalesOrder.objects.filter(
                customer_id=customer_id
            ).values_list('id', flat=True)
            # Find DC IDs that have lines linked to those SOs
            customer_dc_ids = DCLine.objects.filter(
                linked_so_line__so_id__in=customer_so_ids
            ).values_list('dc_id', flat=True).distinct()
            qs = qs.filter(id__in=customer_dc_ids)

        available = qs.values('id', 'dc_no', 'invoice_no', 'warehouse__name')
        return Response([
            {'id': str(dc['id']), 'dc_no': dc['dc_no'], 'invoice_no': dc['invoice_no'] or '', 'warehouse_name': dc['warehouse__name'] or ''}
            for dc in available
        ])


class FreightAdviceOutboundViewSet(viewsets.ModelViewSet):
    """
    Outward Freight ViewSet.
    Full CRUD + payment management + attachments.
    """

    queryset = FreightAdviceOutbound.objects.all()
    permission_classes = [permissions.IsAuthenticated, HasModulePermission]
    module_name = 'Freight Advice'
    filterset_fields = ['transporter', 'status']
    search_fields = ['advice_no', 'customer_name', 'lorry_no']
    ordering_fields = ['created_date', 'freight_date', 'status', 'payable_amount']
    ordering = ['-created_date']

    def get_queryset(self):
        return (
            FreightAdviceOutbound.objects
            .select_related('transporter', 'created_by')
            .prefetch_related('dc_links__dc', 'payments', 'attachments')
            .filter(is_active=True)
        )

    def get_serializer_class(self):
        if self.action in ('create', 'update', 'partial_update'):
            return CreateUpdateFreightSerializer
        elif self.action == 'retrieve':
            return FreightAdviceOutboundDetailSerializer
        return FreightAdviceOutboundListSerializer

    def perform_destroy(self, instance):
        """Delete Outward Freight and rollback: update linked Freight Detail paid/balance, delete payments."""
        from django.db import transaction as db_transaction
        with db_transaction.atomic():
            # Delete all payment records for this freight
            instance.payments.all().delete()
            # If linked to a Freight Detail, reset its paid amount
            if instance.freight_detail:
                fd = instance.freight_detail
                # Recalculate from remaining active outward freights (excluding this one)
                remaining_paid = FreightAdviceOutbound.objects.filter(
                    freight_detail=fd, is_active=True
                ).exclude(id=instance.id).aggregate(total=Sum('total_paid'))['total'] or Decimal('0')
                fd.freight_paid = remaining_paid
                fd.balance_freight = max(Decimal('0'), (fd.total_freight or Decimal('0')) - remaining_paid)
                if fd.balance_freight <= 0 and remaining_paid > 0:
                    fd.status = 'COMPLETED'
                elif remaining_paid > 0:
                    fd.status = 'IN_PROGRESS'
                else:
                    fd.status = 'PENDING'
                fd.save(update_fields=['freight_paid', 'balance_freight', 'status', 'updated_at'])
            # Soft delete
            instance.is_active = False
            instance.save(update_fields=['is_active', 'updated_at'])

    # --- Payment CRUD ---
    @action(detail=True, methods=['post'], url_path='add-payment')
    def add_payment(self, request, pk=None):
        """Add a payment entry to this freight."""
        from .models import FreightPayment
        freight = self.get_object()
        serializer = FreightPaymentSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)

        amount = serializer.validated_data['amount_paid']
        if amount > freight.balance:
            raise ValidationError(
                f'Payment amount ({amount}) exceeds balance ({freight.balance}).'
            )

        payment = FreightPayment.objects.create(
            freight=freight,
            created_by=request.user,
            **serializer.validated_data
        )
        freight.update_payment_status()
        # Sync payment to parent FreightDetail
        self._sync_freight_detail_paid(freight)
        return Response(
            FreightPaymentSerializer(payment).data,
            status=status.HTTP_201_CREATED
        )

    @action(detail=True, methods=['get'], url_path='payments')
    def list_payments(self, request, pk=None):
        """List all payments for this freight."""
        freight = self.get_object()
        payments = freight.payments.all()
        return Response(FreightPaymentSerializer(payments, many=True).data)

    @action(detail=True, methods=['delete'], url_path='delete-payment/(?P<payment_id>[^/.]+)')
    def delete_payment(self, request, pk=None, payment_id=None):
        """Delete a payment entry and recalculate."""
        from .models import FreightPayment
        freight = self.get_object()
        try:
            payment = freight.payments.get(id=payment_id)
        except FreightPayment.DoesNotExist:
            raise ValidationError('Payment not found.')
        payment.delete()
        freight.update_payment_status()
        # Sync payment to parent FreightDetail
        self._sync_freight_detail_paid(freight)
        return Response({'detail': 'Payment deleted'}, status=status.HTTP_200_OK)

    def _sync_freight_detail_paid(self, freight):
        """Sync total paid from all linked Outward Freights back to the parent FreightDetail."""
        if not freight.freight_detail:
            return
        fd = freight.freight_detail
        # Sum total_paid from ALL active outward freights linked to this freight detail
        total = FreightAdviceOutbound.objects.filter(
            freight_detail=fd, is_active=True
        ).aggregate(total=Sum('total_paid'))['total'] or Decimal('0')
        fd.freight_paid = total
        fd.balance_freight = max(Decimal('0'), (fd.total_freight or Decimal('0')) - total)
        if fd.balance_freight <= 0 and total > 0:
            fd.status = 'COMPLETED'
        elif total > 0:
            fd.status = 'IN_PROGRESS'
        else:
            fd.status = 'PENDING'
        fd.save(update_fields=['freight_paid', 'balance_freight', 'status', 'updated_at'])

    # --- Attachment CRUD ---
    @action(detail=True, methods=['post'], url_path='add-attachment')
    def add_attachment(self, request, pk=None):
        """Upload an attachment to this freight."""
        from .models import FreightAttachment
        freight = self.get_object()
        serializer = FreightAttachmentSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        attachment = FreightAttachment.objects.create(
            freight=freight, **serializer.validated_data
        )
        return Response(
            FreightAttachmentSerializer(attachment).data,
            status=status.HTTP_201_CREATED
        )

    @action(detail=True, methods=['delete'], url_path='delete-attachment/(?P<attachment_id>[^/.]+)')
    def delete_attachment(self, request, pk=None, attachment_id=None):
        """Delete an attachment."""
        from .models import FreightAttachment
        freight = self.get_object()
        try:
            att = freight.attachments.get(id=attachment_id)
        except FreightAttachment.DoesNotExist:
            raise ValidationError('Attachment not found.')
        att.delete()
        return Response({'detail': 'Attachment deleted'}, status=status.HTTP_200_OK)

    # --- Status ---
    @action(detail=True, methods=['post'])
    def cancel(self, request, pk=None):
        """Cancel freight advice."""
        freight = self.get_object()
        freight.status = 'CANCELLED'
        freight.save(update_fields=['status', 'updated_at'])
        return Response({'detail': 'Freight cancelled'}, status=status.HTTP_200_OK)


class FreightPaymentViewSet(viewsets.ModelViewSet):
    """Standalone ViewSet for listing/creating/deleting freight payments."""
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        return FreightPayment.objects.select_related(
            'freight__transporter', 'created_by'
        ).filter(freight__is_active=True).order_by('-payment_date')

    def get_serializer_class(self):
        if self.action == 'list':
            return FreightPaymentListSerializer
        return FreightPaymentListSerializer

    def perform_create(self, serializer):
        from django.db import transaction as db_transaction
        with db_transaction.atomic():
            freight = serializer.validated_data['freight']
            amount = serializer.validated_data['amount_paid']
            if amount > freight.balance:
                raise ValidationError(
                    f'Payment amount ({amount}) exceeds balance ({freight.balance}).'
                )
            payment = serializer.save(created_by=self.request.user)
            freight.update_payment_status()
            # Sync to FreightDetail
            self._sync_freight_detail_paid(freight)

    def perform_destroy(self, instance):
        from django.db import transaction as db_transaction
        with db_transaction.atomic():
            freight = instance.freight
            instance.delete()
            freight.update_payment_status()
            self._sync_freight_detail_paid(freight)

    def _sync_freight_detail_paid(self, freight):
        if not freight.freight_detail:
            return
        fd = freight.freight_detail
        total = FreightAdviceOutbound.objects.filter(
            freight_detail=fd, is_active=True
        ).aggregate(total=Sum('total_paid'))['total'] or Decimal('0')
        fd.freight_paid = total
        fd.balance_freight = max(Decimal('0'), (fd.total_freight or Decimal('0')) - total)
        if fd.balance_freight <= 0 and total > 0:
            fd.status = 'COMPLETED'
        elif total > 0:
            fd.status = 'IN_PROGRESS'
        else:
            fd.status = 'PENDING'
        fd.save(update_fields=['freight_paid', 'balance_freight', 'status', 'updated_at'])

    @action(detail=False, methods=['get'], url_path='statement')
    def statement(self, request):
        """Generate freight payment statement with filters. Returns JSON or Word doc."""
        from datetime import date, timedelta
        import io

        # Filters
        period = request.query_params.get('period', '1M')
        freight_id = request.query_params.get('freight_id')
        transporter = request.query_params.get('transporter')
        date_from = request.query_params.get('date_from')
        date_to = request.query_params.get('date_to')
        fmt = request.query_params.get('export', 'json')  # json or docx

        today = date.today()
        if date_from:
            cutoff_from = date.fromisoformat(date_from)
        else:
            period_map = {
                '1W': today - timedelta(weeks=1),
                '2W': today - timedelta(weeks=2),
                '1M': today - timedelta(days=30),
                '3M': today - timedelta(days=90),
                '6M': today - timedelta(days=180),
                '1Y': today - timedelta(days=365),
                '2Y': today - timedelta(days=730),
                'ALL': date(2000, 1, 1),
            }
            cutoff_from = period_map.get(period, today - timedelta(days=30))

        cutoff_to = date.fromisoformat(date_to) if date_to else today

        # Build queryset
        qs = FreightPayment.objects.select_related('freight', 'created_by').filter(
            freight__is_active=True,
            payment_date__gte=cutoff_from,
            payment_date__lte=cutoff_to,
        )
        if freight_id:
            qs = qs.filter(freight_id=freight_id)
        if transporter:
            qs = qs.filter(freight__transporter__name__icontains=transporter)

        payments = qs.order_by('payment_date')

        # Also get outward freight summary for context
        freight_ids = set(payments.values_list('freight_id', flat=True))
        freights = FreightAdviceOutbound.objects.filter(id__in=freight_ids, is_active=True)

        # Build statement rows
        rows = []
        running_total = Decimal('0')
        for p in payments:
            running_total += p.amount_paid
            rows.append({
                'date': str(p.payment_date),
                'advice_no': p.freight.advice_no if p.freight else '',
                'customer': p.freight.customer_name if p.freight else '',
                'transporter': p.freight.transporter.name if p.freight and p.freight.transporter else '',
                'lorry_no': p.freight.lorry_no if p.freight else '',
                'destination': p.freight.destination if p.freight else '',
                'amount': str(p.amount_paid),
                'mode': p.get_payment_mode_display() if hasattr(p, 'get_payment_mode_display') else p.payment_mode,
                'reference': p.reference_no or '',
                'remarks': p.remarks or '',
                'running_total': str(running_total),
            })

        total_payable = sum(f.payable_amount or 0 for f in freights)
        total_paid = sum(f.total_paid or 0 for f in freights)
        total_balance = sum(f.balance or 0 for f in freights)

        summary = {
            'period_from': str(cutoff_from),
            'period_to': str(cutoff_to),
            'total_payments': len(rows),
            'total_amount': str(running_total),
            'total_payable': str(total_payable),
            'total_paid_all': str(total_paid),
            'total_balance': str(total_balance),
            'freight_count': len(freight_ids),
        }

        if fmt == 'docx':
            return self._generate_word_statement(rows, summary, request)

        return Response({'summary': summary, 'rows': rows})

    def _generate_word_statement(self, rows, summary, request):
        """Generate advanced Word document statement with full details."""
        import io
        from datetime import date as dt_date
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from django.http import HttpResponse

        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(9)

        for section in doc.sections:
            section.top_margin = Cm(1.2)
            section.bottom_margin = Cm(1.2)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
            section.page_width = Cm(29.7)  # A4 Landscape
            section.page_height = Cm(21)

        def set_cell_bg(cell, color):
            shading = cell._element.get_or_add_tcPr()
            shd = shading.makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'})
            shading.append(shd)

        def add_styled_cell(cell, text, bold=False, size=8, color='000000', align='LEFT', bg=None):
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = {'LEFT': WD_ALIGN_PARAGRAPH.LEFT, 'CENTER': WD_ALIGN_PARAGRAPH.CENTER, 'RIGHT': WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.LEFT)
            run = p.add_run(str(text))
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.color.rgb = RGBColor.from_string(color)
            run.font.name = 'Calibri'
            if bg:
                set_cell_bg(cell, bg)

        today = dt_date.today().strftime('%d %b %Y')
        period_from = summary.get('period_from', '')
        period_to = summary.get('period_to', '')

        # ═══ HEADER BANNER ═══
        h_table = doc.add_table(rows=1, cols=2)
        h_table.columns[0].width = Cm(16)
        h_table.columns[1].width = Cm(10)
        left = h_table.rows[0].cells[0]
        right = h_table.rows[0].cells[1]
        set_cell_bg(left, '1E3A8A')
        set_cell_bg(right, '1E3A8A')

        p = left.paragraphs[0]
        run = p.add_run('LANCER ERP')
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        p.add_run('\n')
        run2 = p.add_run('Freight Payment Statement')
        run2.font.size = Pt(11)
        run2.font.color.rgb = RGBColor(191, 219, 254)

        p2 = right.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run3 = p2.add_run(f'Report Date: {today}')
        run3.font.size = Pt(9)
        run3.font.color.rgb = RGBColor(191, 219, 254)
        p2.add_run('\n')
        run4 = p2.add_run(f'Period: {period_from} to {period_to}')
        run4.font.size = Pt(9)
        run4.font.color.rgb = RGBColor(255, 255, 255)
        run4.font.bold = True

        doc.add_paragraph()

        # ═══ SUMMARY CARDS (4 columns) ═══
        s_table = doc.add_table(rows=2, cols=4)
        s_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        labels = ['Total Payments', 'Total Amount Paid', 'Total Freight Payable', 'Outstanding Balance']
        values = [
            str(summary['total_payments']),
            f"₹{float(summary['total_amount']):,.2f}",
            f"₹{float(summary['total_payable']):,.2f}",
            f"₹{float(summary['total_balance']):,.2f}",
        ]
        colors = ['1E40AF', '059669', '0891B2', 'D97706' if float(summary['total_balance']) > 0 else '059669']
        bgs = ['EFF6FF', 'ECFDF5', 'ECFEFF', 'FFFBEB' if float(summary['total_balance']) > 0 else 'ECFDF5']
        for i in range(4):
            add_styled_cell(s_table.rows[0].cells[i], labels[i], size=7, color='64748B', bold=True, bg=bgs[i])
            add_styled_cell(s_table.rows[1].cells[i], values[i], size=14, color=colors[i], bold=True, bg=bgs[i])

        doc.add_paragraph()

        # ═══ ADDITIONAL INFO ═══
        info_table = doc.add_table(rows=1, cols=4)
        info_data = [
            ('Freight Count', str(summary.get('freight_count', 0))),
            ('Total Paid (All)', f"₹{float(summary.get('total_paid_all', 0)):,.2f}"),
            ('Period From', period_from),
            ('Period To', period_to),
        ]
        for i, (label, val) in enumerate(info_data):
            cell = info_table.rows[0].cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            r1 = p.add_run(f'{label}: ')
            r1.font.size = Pt(8)
            r1.font.color.rgb = RGBColor(100, 116, 139)
            r2 = p.add_run(val)
            r2.font.size = Pt(8)
            r2.font.bold = True
            r2.font.color.rgb = RGBColor(30, 41, 59)

        doc.add_paragraph()

        # ═══ SECTION: PAYMENT LEDGER ═══
        sh = doc.add_paragraph()
        run = sh.add_run('PAYMENT LEDGER')
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(30, 58, 138)

        if not rows:
            p = doc.add_paragraph()
            run = p.add_run('No payments found for the selected period.')
            run.font.size = Pt(10)
            run.font.italic = True
            run.font.color.rgb = RGBColor(148, 163, 184)
        else:
            headers = ['#', 'Date', 'Freight No', 'Transporter', 'Lorry No', 'Destination', 'Amount (₹)', 'Mode', 'Reference', 'Remarks', 'Running Total (₹)']
            table = doc.add_table(rows=1, cols=len(headers))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Header
            for j, header in enumerate(headers):
                add_styled_cell(table.rows[0].cells[j], header, bold=True, size=7, color='FFFFFF', align='CENTER', bg='1E3A8A')

            # Data rows
            for i, row in enumerate(rows):
                r = table.add_row()
                bg = 'F8FAFC' if i % 2 == 0 else 'FFFFFF'
                data = [
                    str(i + 1), row['date'], row['advice_no'], row['transporter'],
                    row.get('lorry_no', ''), row.get('destination', ''),
                    f"₹{float(row['amount']):,.2f}", row['mode'],
                    row.get('reference', ''), row.get('remarks', ''),
                    f"₹{float(row['running_total']):,.2f}",
                ]
                aligns = ['CENTER', 'LEFT', 'LEFT', 'LEFT', 'LEFT', 'LEFT', 'RIGHT', 'CENTER', 'LEFT', 'LEFT', 'RIGHT']
                for j, (val, align) in enumerate(zip(data, aligns)):
                    c = '059669' if j == 6 else ('1E293B' if j == 10 else '334155')
                    b = j == 10
                    add_styled_cell(r.cells[j], val, size=7, color=c, align=align, bg=bg, bold=b)

            # Total row
            tr = table.add_row()
            for j in range(len(headers)):
                set_cell_bg(tr.cells[j], 'E2E8F0')
            add_styled_cell(tr.cells[0], '', bg='E2E8F0')
            add_styled_cell(tr.cells[5], 'TOTAL', bold=True, size=8, color='1E293B', align='RIGHT', bg='E2E8F0')
            add_styled_cell(tr.cells[6], f"₹{float(summary['total_amount']):,.2f}", bold=True, size=9, color='059669', align='RIGHT', bg='E2E8F0')
            add_styled_cell(tr.cells[10], f"₹{float(summary['total_amount']):,.2f}", bold=True, size=9, color='1E293B', align='RIGHT', bg='E2E8F0')

        doc.add_paragraph()

        # ═══ TRANSPORTER BREAKDOWN ═══
        transporters = {}
        for row in rows:
            t = row.get('transporter') or 'Unknown'
            if t not in transporters:
                transporters[t] = {'count': 0, 'total': 0}
            transporters[t]['count'] += 1
            transporters[t]['total'] += float(row['amount'])

        if transporters:
            bh = doc.add_paragraph()
            run = bh.add_run('TRANSPORTER-WISE BREAKDOWN')
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(30, 58, 138)

            bt = doc.add_table(rows=1, cols=3)
            bt.alignment = WD_TABLE_ALIGNMENT.CENTER
            for j, h in enumerate(['Transporter', 'Payments', 'Total Amount']):
                add_styled_cell(bt.rows[0].cells[j], h, bold=True, size=8, color='FFFFFF', align='CENTER', bg='0891B2')

            for i, (name, data) in enumerate(sorted(transporters.items(), key=lambda x: -x[1]['total'])):
                r = bt.add_row()
                bg = 'F0FDFA' if i % 2 == 0 else 'FFFFFF'
                add_styled_cell(r.cells[0], name, size=8, color='334155', bg=bg, bold=True)
                add_styled_cell(r.cells[1], str(data['count']), size=8, color='334155', align='CENTER', bg=bg)
                add_styled_cell(r.cells[2], f"₹{data['total']:,.2f}", size=8, color='059669', align='RIGHT', bg=bg, bold=True)

        doc.add_paragraph()

        # ═══ PAYMENT MODE BREAKDOWN ═══
        modes = {}
        for row in rows:
            m = row.get('mode') or 'Other'
            if m not in modes:
                modes[m] = {'count': 0, 'total': 0}
            modes[m]['count'] += 1
            modes[m]['total'] += float(row['amount'])

        if modes:
            mh = doc.add_paragraph()
            run = mh.add_run('PAYMENT MODE BREAKDOWN')
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(30, 58, 138)

            mt = doc.add_table(rows=1, cols=3)
            mt.alignment = WD_TABLE_ALIGNMENT.CENTER
            for j, h in enumerate(['Payment Mode', 'Count', 'Total Amount']):
                add_styled_cell(mt.rows[0].cells[j], h, bold=True, size=8, color='FFFFFF', align='CENTER', bg='7C3AED')

            for i, (mode, data) in enumerate(sorted(modes.items(), key=lambda x: -x[1]['total'])):
                r = mt.add_row()
                bg = 'F5F3FF' if i % 2 == 0 else 'FFFFFF'
                add_styled_cell(r.cells[0], mode, size=8, color='334155', bg=bg, bold=True)
                add_styled_cell(r.cells[1], str(data['count']), size=8, color='334155', align='CENTER', bg=bg)
                add_styled_cell(r.cells[2], f"₹{data['total']:,.2f}", size=8, color='7C3AED', align='RIGHT', bg=bg, bold=True)

        doc.add_paragraph()
        doc.add_paragraph()

        # ═══ FOOTER ═══
        ft = doc.add_table(rows=1, cols=2)
        set_cell_bg(ft.rows[0].cells[0], 'F1F5F9')
        set_cell_bg(ft.rows[0].cells[1], 'F1F5F9')
        p1 = ft.rows[0].cells[0].paragraphs[0]
        r1 = p1.add_run(f'Generated on {today}')
        r1.font.size = Pt(7)
        r1.font.color.rgb = RGBColor(148, 163, 184)
        p2 = ft.rows[0].cells[1].paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r2 = p2.add_run('Lancer ERP — Freight Payment Statement')
        r2.font.size = Pt(7)
        r2.font.color.rgb = RGBColor(148, 163, 184)
        r2.font.italic = True

        # Save
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="freight-payment-statement-{summary["period_from"]}-to-{summary["period_to"]}.docx"'
        return response


class ReceivableLedgerViewSet(viewsets.ModelViewSet):
    """
    ViewSet for accounts receivable ledger.
    Tracks invoices, payments, and aging.
    """

    queryset = ReceivableLedger.objects.all()
    permission_classes = [permissions.IsAuthenticated, HasModulePermission]
    module_name = 'Receivable'
    filterset_fields = ['customer', 'payment_status', 'escalation_flag']
    search_fields = ['customer__customer_name', 'invoice_reference__invoice_no']
    ordering_fields = ['due_date', 'amount']
    ordering = ['-invoice_date']

    def get_queryset(self):
        return (
            ReceivableLedger.objects
            .select_related('customer', 'invoice_reference')
            .prefetch_related('reminders')
            .filter(is_active=True)
        )

    def get_serializer_class(self):
        if self.action in ('create', 'update', 'partial_update'):
            return CreateUpdateReceivableSerializer
        if self.action == 'retrieve':
            return ReceivableLedgerDetailSerializer
        return ReceivableLedgerListSerializer

    def perform_destroy(self, instance):
        instance.is_active = False
        instance.save(update_fields=['is_active', 'updated_at'])

    @action(detail=True, methods=['post'])
    def record_payment(self, request, pk=None):
        """Record payment received"""
        receivable = self.get_object()

        from decimal import Decimal

        amount = Decimal(request.data.get('amount', 0))
        if amount <= 0:
            raise ValidationError("Amount must be positive")

        ReceivableService.record_payment_received(receivable, amount)

        return Response(
            {
                'detail': 'Payment recorded',
                'balance': str(receivable.balance),
            },
            status=status.HTTP_200_OK
        )

    @action(detail=False, methods=['get'])
    def overdue(self, request):
        """Get overdue receivables"""
        days_overdue = request.query_params.get('days_overdue', 0, type=int)

        receivables = ReceivableSelector.get_overdue_receivables(days_overdue)
        serializer = ReceivableLedgerListSerializer(receivables, many=True)

        return Response(serializer.data)

    @action(detail=False, methods=['get'])
    def aging_summary(self, request):
        """Get aging bucket summary"""
        aging_data = ReceivableSelector.get_aging_summary()

        return Response(aging_data)

    @action(detail=False, methods=['get'])
    def customer_summary(self, request):
        """Get reconciliation summary for customer"""
        customer_id = request.query_params.get('customer_id')

        if not customer_id:
            raise ValidationError("customer_id query parameter required")

        summary = SalesReconciliationSelector.get_reconciliation_summary_by_customer(
            customer_id
        )

        return Response(summary)


class SalesReconciliationViewSet(viewsets.ViewSet):
    """
    ViewSet for sales reconciliation.
    Provides PO → SO → DC → Invoice → Payment trails.
    """

    permission_classes = [permissions.IsAuthenticated]

    @action(detail=False, methods=['get'])
    def sales_order_trail(self, request):
        """Get complete reconciliation trail for sales order"""
        so_id = request.query_params.get('so_id')

        if not so_id:
            raise ValidationError("so_id query parameter required")

        trail = SalesReconciliationSelector.get_sales_reconciliation_trail(so_id)

        return Response(trail)

    @action(detail=False, methods=['get'])
    def invoice_matching(self, request):
        """Get invoice to receivable matching details"""
        invoice_check_id = request.query_params.get('invoice_check_id')

        if not invoice_check_id:
            raise ValidationError("invoice_check_id query parameter required")

        matching = SalesReconciliationSelector.get_invoice_to_receivable_matching(
            invoice_check_id
        )

        return Response(matching)
