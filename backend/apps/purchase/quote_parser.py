"""
Enhanced quote document parser with confidence scoring and DB matching.

Supports: PDF, images (JPG/PNG), Word (DOCX), Excel (XLSX/XLS), CSV, TXT
Falls back gracefully when optional dependencies are unavailable.
"""

import re
import csv
import json
import logging
from io import BytesIO, StringIO
from decimal import Decimal, InvalidOperation
from datetime import datetime, timedelta

logger = logging.getLogger(__name__)


class ExtractedField:
    """Wrapper for an extracted field with confidence and source context."""

    __slots__ = ('value', 'confidence', 'source')

    def __init__(self, value, confidence=0.0, source=''):
        self.value = value
        self.confidence = round(min(max(confidence, 0.0), 1.0), 2)
        self.source = source

    def to_dict(self):
        d = {'value': self.value, 'confidence': self.confidence}
        if self.source:
            d['source'] = self.source
        return d

    def __bool__(self):
        return self.value is not None and self.value != ''


class QuoteParser:
    """Enhanced quote document parser with confidence scoring."""

    SUPPORTED_FORMATS = ['pdf', 'jpg', 'jpeg', 'png', 'doc', 'docx', 'xlsx', 'xls', 'csv', 'txt']

    # Date patterns for normalization
    DATE_PATTERNS = [
        (r'(\d{2})[/-](\d{2})[/-](\d{4})', 'DMY'),   # DD-MM-YYYY or DD/MM/YYYY
        (r'(\d{4})[/-](\d{2})[/-](\d{2})', 'YMD'),   # YYYY-MM-DD
        (r'(\d{1,2})\s+(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\s+(\d{4})', 'DMonY'),
        (r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\s+(\d{1,2}),?\s+(\d{4})', 'MonDY'),
    ]

    MONTH_MAP = {
        'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
        'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12,
    }

    PAYMENT_TERMS_MAP = {
        'advance': 'ADVANCE',
        'immediate': 'IMMEDIATE',
        'cod': 'COD',
        'cash on delivery': 'COD',
        'net 7': 'NET_7',
        'net 15': 'NET_15',
        'net 30': 'NET_30',
        'net 45': 'NET_45',
        'net 60': 'NET_60',
        'net 90': 'NET_90',
        '7 days': 'NET_7',
        '15 days': 'NET_15',
        '30 days': 'NET_30',
        '45 days': 'NET_45',
        '60 days': 'NET_60',
        '90 days': 'NET_90',
        'against delivery': 'AGAINST_DELIVERY',
        'against proforma': 'AGAINST_PI',
        'lc': 'LC',
        'letter of credit': 'LC',
    }

    FREIGHT_TERMS_MAP = {
        'paid': 'PAID',
        'to pay': 'TO_PAY',
        'to-pay': 'TO_PAY',
        'topay': 'TO_PAY',
        'fob': 'FOB',
        'cif': 'CIF',
        'ex-works': 'EX_WORKS',
        'ex works': 'EX_WORKS',
        'exworks': 'EX_WORKS',
        'door delivery': 'DOOR_DELIVERY',
        'franco': 'FRANCO',
        'included': 'INCLUDED',
        'extra': 'EXTRA',
    }

    # ------------------------------------------------------------------ #
    #  Public API                                                         #
    # ------------------------------------------------------------------ #

    @classmethod
    def parse_raw_text(cls, text):
        """Parse raw text input (WhatsApp, email, OCR, chat).

        Returns the same structure as file-based parsing, plus a
        structured_output key with the document extraction schema.
        """
        if not text or not text.strip():
            return {
                'status': 'failed',
                'reason': 'Empty input text',
            }

        result = cls._build_result(text.strip(), fmt='text', pages=1)

        # Build the structured document extraction schema
        extracted = result.get('extracted', {})
        detailed = result.get('extracted_detailed', {})
        line_items = result.get('line_items', [])

        def _field(key, fallback_key=None):
            d = detailed.get(key) or detailed.get(fallback_key or '', {})
            if isinstance(d, dict):
                return {'value': d.get('value'), 'confidence': d.get('confidence', 0)}
            return {'value': None, 'confidence': 0}

        # Detect document type
        text_lower = text.lower()
        if any(w in text_lower for w in ['invoice', 'inv no', 'bill']):
            doc_type = {'value': 'invoice', 'confidence': 0.85}
        elif any(w in text_lower for w in ['quotation', 'quote', 'offer', 'proforma']):
            doc_type = {'value': 'quotation', 'confidence': 0.85}
        elif any(w in text_lower for w in ['rfq', 'request for']):
            doc_type = {'value': 'RFQ', 'confidence': 0.85}
        elif any(w in text_lower for w in ['ledger', 'statement', 'balance']):
            doc_type = {'value': 'ledger', 'confidence': 0.7}
        else:
            doc_type = {'value': 'unknown', 'confidence': 0.4}

        # Extract phone
        phone_match = re.search(r'(?<!\d)([6-9]\d{9})(?!\d)', text)
        phone = {'value': phone_match.group(1) if phone_match else None,
                 'confidence': 0.95 if phone_match else 0}

        # Extract email
        email_match = re.search(r'[\w.+-]+@[\w-]+\.[\w.-]+', text)
        email = {'value': email_match.group(0) if email_match else None,
                 'confidence': 0.95 if email_match else 0}

        # Build products
        products = []
        for item in line_items:
            p = {
                'name': {'value': (item.get('product_name') or {}).get('value') or
                                  (item.get('description') or {}).get('value'),
                         'confidence': (item.get('product_name') or {}).get('confidence', 0.5)},
                'quantity': {'value': (item.get('quantity') or item.get('qty') or {}).get('value'),
                             'confidence': (item.get('quantity') or item.get('qty') or {}).get('confidence', 0.5)},
                'unit': {'value': (item.get('uom') or item.get('unit') or {}).get('value'),
                         'confidence': (item.get('uom') or item.get('unit') or {}).get('confidence', 0.5)},
                'price_per_unit': {'value': (item.get('unit_price') or item.get('rate') or {}).get('value'),
                                   'confidence': (item.get('unit_price') or item.get('rate') or {}).get('confidence', 0.5)},
                'total_price': {'value': (item.get('total') or item.get('amount') or {}).get('value'),
                                'confidence': (item.get('total') or item.get('amount') or {}).get('confidence', 0.5)},
            }
            if item.get('matched_product_id'):
                p['matched_product_id'] = item['matched_product_id']
            products.append(p)

        # Validation
        validation_flags = []
        for i, p in enumerate(products):
            qty = cls._safe_float(p['quantity']['value'])
            price = cls._safe_float(p['price_per_unit']['value'])
            total = cls._safe_float(p['total_price']['value'])
            if qty and price and total:
                expected = qty * price
                if abs(expected - total) > 1:
                    validation_flags.append(
                        f'Line {i+1}: qty({qty}) x price({price}) = {expected}, but total shows {total}')

        # Confidence summary
        all_confs = [doc_type['confidence'], phone['confidence'], email['confidence']]
        for p in products:
            all_confs.extend([p['name']['confidence'], p['quantity']['confidence'],
                              p['price_per_unit']['confidence']])
        valid_confs = [c for c in all_confs if c > 0]
        overall = round(sum(valid_confs) / max(len(valid_confs), 1), 2)
        low_fields = []
        field_map = {'document_type': doc_type, 'phone': phone, 'email': email,
                     'vendor_name': _field('vendor_name')}
        for fname, fval in field_map.items():
            if fval['confidence'] < 0.5 and fval['value']:
                low_fields.append(fname)

        subtotal = _field('subtotal', 'total_amount')
        tax = _field('gst_amount', 'tax_amount')
        grand_total = _field('grand_total', 'total_amount')

        structured = {
            'status': 'success',
            'data': {
                'document_type': doc_type,
                'vendor_name': _field('vendor_name'),
                'customer_name': {'value': None, 'confidence': 0},
                'phone': phone,
                'email': email,
                'date': _field('quote_date'),
                'products': products,
                'subtotal': subtotal,
                'tax': tax,
                'grand_total': grand_total,
                'currency': {'value': 'INR', 'confidence': 1.0},
                'notes': _field('notes', 'remarks'),
            },
            'confidence_summary': {
                'overall_confidence': overall,
                'low_confidence_fields': low_fields,
            },
            'validation_flags': validation_flags,
            # Also include the raw parser output for advanced use
            'raw_parser': {
                'raw_text': text[:5000],  # first 5000 chars of raw text
                'extracted': extracted,
                'line_items': line_items,
                'warnings': result.get('warnings', []),
                'confidence': result.get('confidence', 0),
            },
        }
        return structured

    @staticmethod
    def _safe_float(val):
        if val is None:
            return None
        try:
            return float(str(val).replace(',', ''))
        except (ValueError, TypeError):
            return None

    @classmethod
    def parse(cls, file_obj, filename):
        """Main entry point - detect format and parse.

        Args:
            file_obj: A file-like object (e.g. Django UploadedFile).
            filename: Original filename with extension.

        Returns:
            dict with keys: success, raw_text, extracted, extracted_detailed,
            line_items, confidence, warnings, format, pages, and optionally error.
        """
        ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''

        if ext not in cls.SUPPORTED_FORMATS:
            return {
                'success': False,
                'error': f'Unsupported format: {ext}',
                'extracted': {},
                'extracted_detailed': {},
                'line_items': [],
                'warnings': [],
            }

        try:
            if ext == 'pdf':
                return cls._parse_pdf(file_obj)
            elif ext in ('jpg', 'jpeg', 'png'):
                return cls._parse_image(file_obj)
            elif ext in ('doc', 'docx'):
                return cls._parse_docx(file_obj)
            elif ext in ('xlsx', 'xls'):
                return cls._parse_excel(file_obj)
            elif ext == 'csv':
                return cls._parse_csv(file_obj)
            elif ext == 'txt':
                return cls._parse_txt(file_obj)
        except Exception as e:
            logger.error(f'Quote parsing error: {e}', exc_info=True)
            return {
                'success': False,
                'error': str(e),
                'extracted': {},
                'extracted_detailed': {},
                'line_items': [],
                'warnings': [],
            }

    # ------------------------------------------------------------------ #
    #  Date Normalisation                                                 #
    # ------------------------------------------------------------------ #

    @classmethod
    def normalize_date(cls, date_str):
        """Convert any recognized date format to YYYY-MM-DD string."""
        if not date_str:
            return None
        for pattern, fmt in cls.DATE_PATTERNS:
            m = re.search(pattern, str(date_str), re.IGNORECASE)
            if m:
                groups = m.groups()
                try:
                    if fmt == 'DMY':
                        d, mo, y = int(groups[0]), int(groups[1]), int(groups[2])
                        return f"{y}-{str(mo).zfill(2)}-{str(d).zfill(2)}"
                    elif fmt == 'YMD':
                        return f"{groups[0]}-{groups[1].zfill(2)}-{groups[2].zfill(2)}"
                    elif fmt == 'DMonY':
                        month = cls.MONTH_MAP.get(groups[1][:3].lower(), 1)
                        return f"{groups[2]}-{str(month).zfill(2)}-{groups[0].zfill(2)}"
                    elif fmt == 'MonDY':
                        month = cls.MONTH_MAP.get(groups[0][:3].lower(), 1)
                        return f"{groups[2]}-{str(month).zfill(2)}-{groups[1].zfill(2)}"
                except (ValueError, IndexError):
                    pass
        return None

    # ------------------------------------------------------------------ #
    #  DB Matching Helpers                                                #
    # ------------------------------------------------------------------ #

    @classmethod
    def match_vendor(cls, vendor_name):
        """Try to match vendor name against DB. Returns (vendor_id, confidence)."""
        if not vendor_name:
            return None, 0.0
        try:
            from master.models import Vendor
            # Exact match
            vendor = Vendor.objects.filter(vendor_name__iexact=vendor_name.strip()).first()
            if vendor:
                return str(vendor.id), 0.95
            # Partial / contains match
            short = vendor_name.strip()[:20]
            vendor = Vendor.objects.filter(vendor_name__icontains=short).first()
            if vendor:
                return str(vendor.id), 0.75
        except Exception:
            pass
        return None, 0.0

    @classmethod
    def match_rfq(cls, rfq_no):
        """Try to match RFQ number against DB. Returns (rfq_id, confidence)."""
        if not rfq_no:
            return None, 0.0
        try:
            from purchase.models import RFQHeader
            rfq = RFQHeader.objects.filter(rfq_no__iexact=rfq_no.strip()).first()
            if rfq:
                return str(rfq.id), 0.95
            # Partial match — try contains
            rfq = RFQHeader.objects.filter(rfq_no__icontains=rfq_no.strip()).first()
            if rfq:
                return str(rfq.id), 0.70
        except Exception:
            pass
        return None, 0.0

    @classmethod
    def match_product(cls, product_name=None, sku=None):
        """Try to match product by SKU or name. Returns product_id or None."""
        if not product_name and not sku:
            return None
        try:
            from master.models import Product
            if sku:
                p = Product.objects.filter(sku_code__iexact=sku.strip()).first()
                if p:
                    return str(p.id)
            if product_name:
                p = Product.objects.filter(product_name__icontains=product_name.strip()[:20]).first()
                if p:
                    return str(p.id)
        except Exception:
            pass
        return None

    # ------------------------------------------------------------------ #
    #  Format-Specific Parsers                                            #
    # ------------------------------------------------------------------ #

    @classmethod
    def _parse_pdf(cls, file_obj):
        """Extract text from PDF - try text extraction first, fall back to OCR."""
        text = ''
        table_data = []
        pages = 0

        # Try pdfplumber first (text-based PDFs)
        try:
            import pdfplumber
            file_obj.seek(0)
            with pdfplumber.open(file_obj) as pdf:
                pages = len(pdf.pages)
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + '\n'

                    # Extract tables with pdfplumber's table detection
                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            table_data.append(table)
                            for row in table:
                                if row:
                                    text += ' | '.join(str(cell or '') for cell in row) + '\n'
        except ImportError:
            logger.info('pdfplumber not installed, skipping')
        except Exception as e:
            logger.warning(f'pdfplumber failed: {e}')

        # If no text extracted, try PyPDF2
        if not text.strip():
            try:
                import PyPDF2
                file_obj.seek(0)
                reader = PyPDF2.PdfReader(file_obj)
                pages = len(reader.pages)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + '\n'
            except ImportError:
                logger.info('PyPDF2 not installed, skipping')
            except Exception as e:
                logger.warning(f'PyPDF2 failed: {e}')

        # If still no text, try OCR
        if not text.strip():
            try:
                text = cls._ocr_file(file_obj, 'pdf')
            except ImportError as e:
                logger.info(f'OCR not available: {e}')
            except Exception as e:
                logger.warning(f'OCR failed: {e}')

        if not text.strip():
            return {
                'success': False,
                'error': 'Could not extract text from PDF. '
                         'The file may be a scanned image (OCR requires pytesseract).',
                'raw_text': '',
                'extracted': {},
                'extracted_detailed': {},
                'line_items': [],
                'warnings': [],
            }

        return cls._build_result(text, fmt='pdf', pages=pages, table_data=table_data)

    @classmethod
    def _parse_image(cls, file_obj):
        """Extract text from image using OCR."""
        try:
            text = cls._ocr_file(file_obj, 'image')
        except ImportError as e:
            logger.info(f'OCR not available: {e}')
            return {
                'success': False,
                'error': str(e),
                'hint': 'For image extraction, copy the text from the image and paste it in the text input instead.',
                'extracted': {},
                'extracted_detailed': {},
                'line_items': [],
                'warnings': [str(e)],
            }
        except Exception as e:
            return {
                'success': False,
                'error': f'OCR failed: {e}',
                'extracted': {},
                'extracted_detailed': {},
                'line_items': [],
                'warnings': [],
            }

        if not text.strip():
            return {
                'success': False,
                'error': 'Could not extract text from image',
                'raw_text': '',
                'extracted': {},
                'extracted_detailed': {},
                'line_items': [],
                'warnings': [],
            }

        return cls._build_result(text, fmt='image', pages=1)

    @classmethod
    def _parse_docx(cls, file_obj):
        """Extract text from Word document."""
        text = ''
        try:
            from docx import Document
            file_obj.seek(0)
            doc = Document(file_obj)

            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + '\n'

            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    text += row_text + '\n'
        except ImportError:
            return {
                'success': False,
                'error': 'python-docx not installed',
                'extracted': {},
                'extracted_detailed': {},
                'line_items': [],
                'warnings': [],
            }
        except Exception as e:
            return {
                'success': False,
                'error': f'DOCX parsing failed: {e}',
                'extracted': {},
                'extracted_detailed': {},
                'line_items': [],
                'warnings': [],
            }

        if not text.strip():
            return {
                'success': False,
                'error': 'No text content found in document',
                'raw_text': '',
                'extracted': {},
                'extracted_detailed': {},
                'line_items': [],
                'warnings': [],
            }

        return cls._build_result(text, fmt='docx', pages=1)

    @classmethod
    def _parse_excel(cls, file_obj):
        """Extract data from Excel file."""
        text = ''
        try:
            import openpyxl
            file_obj.seek(0)
            wb = openpyxl.load_workbook(file_obj, data_only=True)

            for sheet in wb.sheetnames:
                ws = wb[sheet]
                for row in ws.iter_rows(values_only=True):
                    row_text = ' | '.join(str(cell or '') for cell in row)
                    if row_text.strip().replace('|', '').strip():
                        text += row_text + '\n'
        except ImportError:
            return {
                'success': False,
                'error': 'openpyxl not installed',
                'extracted': {},
                'extracted_detailed': {},
                'line_items': [],
                'warnings': [],
            }
        except Exception as e:
            return {
                'success': False,
                'error': f'Excel parsing failed: {e}',
                'extracted': {},
                'extracted_detailed': {},
                'line_items': [],
                'warnings': [],
            }

        if not text.strip():
            return {
                'success': False,
                'error': 'No data found in Excel file',
                'raw_text': '',
                'extracted': {},
                'extracted_detailed': {},
                'line_items': [],
                'warnings': [],
            }

        return cls._build_result(text, fmt='excel', pages=1)

    @classmethod
    def _parse_csv(cls, file_obj):
        """Extract data from CSV file."""
        text = ''
        try:
            file_obj.seek(0)
            raw = file_obj.read()
            if isinstance(raw, bytes):
                raw = raw.decode('utf-8', errors='replace')
            reader = csv.reader(StringIO(raw))
            for row in reader:
                if row:
                    text += ' | '.join(row) + '\n'
        except Exception as e:
            return {
                'success': False,
                'error': f'CSV parsing failed: {e}',
                'extracted': {},
                'extracted_detailed': {},
                'line_items': [],
                'warnings': [],
            }

        if not text.strip():
            return {
                'success': False,
                'error': 'No data found in CSV file',
                'raw_text': '',
                'extracted': {},
                'extracted_detailed': {},
                'line_items': [],
                'warnings': [],
            }

        return cls._build_result(text, fmt='csv', pages=1)

    @classmethod
    def _parse_txt(cls, file_obj):
        """Extract data from plain text file."""
        try:
            file_obj.seek(0)
            raw = file_obj.read()
            if isinstance(raw, bytes):
                text = raw.decode('utf-8', errors='replace')
            else:
                text = raw
        except Exception as e:
            return {
                'success': False,
                'error': f'TXT parsing failed: {e}',
                'extracted': {},
                'extracted_detailed': {},
                'line_items': [],
                'warnings': [],
            }

        if not text.strip():
            return {
                'success': False,
                'error': 'No data found in text file',
                'raw_text': '',
                'extracted': {},
                'extracted_detailed': {},
                'line_items': [],
                'warnings': [],
            }

        return cls._build_result(text, fmt='txt', pages=1)

    # ------------------------------------------------------------------ #
    #  OCR Helper                                                         #
    # ------------------------------------------------------------------ #

    @classmethod
    def _ocr_file(cls, file_obj, file_type):
        """Perform OCR on file. Tries Tesseract first, falls back to EasyOCR."""
        from PIL import Image

        file_obj.seek(0)

        # Strategy 1: Try Tesseract
        try:
            import pytesseract
            pytesseract.get_tesseract_version()
            logger.info('Using Tesseract OCR')
            if file_type == 'pdf':
                from pdf2image import convert_from_bytes
                images = convert_from_bytes(file_obj.read())
                return '\n'.join(pytesseract.image_to_string(img) for img in images)
            else:
                return pytesseract.image_to_string(Image.open(file_obj))
        except Exception as e:
            logger.info(f'Tesseract not available: {e}')

        # Strategy 2: Try EasyOCR (pure Python, no system binary needed)
        try:
            import easyocr
            import numpy as np
            logger.info('Using EasyOCR fallback (first run may download models)')
            file_obj.seek(0)
            img = Image.open(file_obj).convert('RGB')
            img_array = np.array(img)

            # Cache the reader at class level to avoid re-init
            if not hasattr(cls, '_easyocr_reader') or cls._easyocr_reader is None:
                cls._easyocr_reader = easyocr.Reader(['en'], gpu=False, verbose=False)
            reader = cls._easyocr_reader

            results = reader.readtext(img_array, detail=1)
            # Sort by vertical position then horizontal
            results.sort(key=lambda r: (r[0][0][1], r[0][0][0]))
            lines = []
            current_y = -1
            current_line = []
            for bbox, text_val, conf in results:
                y = bbox[0][1]
                if current_y >= 0 and abs(y - current_y) > 15:
                    lines.append(' '.join(current_line))
                    current_line = []
                current_line.append(text_val)
                current_y = y
            if current_line:
                lines.append(' '.join(current_line))
            return '\n'.join(lines)
        except ImportError:
            logger.info('EasyOCR not installed either')
        except Exception as e:
            logger.warning(f'EasyOCR failed: {e}')

        raise ImportError(
            'No OCR engine available. Install either: '
            'Tesseract OCR (https://github.com/UB-Mannheim/tesseract/wiki) '
            'or run: pip install easyocr'
        )

    # ------------------------------------------------------------------ #
    #  Result Builder                                                     #
    # ------------------------------------------------------------------ #

    @classmethod
    def _build_result(cls, text, fmt='unknown', pages=1, table_data=None):
        """Build the full enhanced result from extracted text.

        Returns the response dict with both backward-compatible `extracted`
        (flat values) and `extracted_detailed` (with confidence scores).
        """
        detailed, warnings = cls._extract_fields_detailed(text, table_data=table_data)
        line_items_detailed = cls._extract_line_items(text, table_data=table_data)

        # --- DB matching ---
        # Vendor
        vendor_name_val = detailed.get('vendor_name', ExtractedField(None)).value
        vendor_id, vendor_conf = cls.match_vendor(vendor_name_val)
        if vendor_id:
            detailed['vendor_id'] = ExtractedField(vendor_id, vendor_conf)
        elif vendor_name_val:
            warnings.append(f'Vendor "{vendor_name_val}" not found in database')

        # Check vendor email matches DB
        vendor_email_val = detailed.get('vendor_email', ExtractedField(None)).value
        if vendor_id and vendor_email_val:
            try:
                from master.models import Vendor
                v = Vendor.objects.filter(id=vendor_id).first()
                if v and v.contact_email and v.contact_email.lower() != vendor_email_val.lower():
                    warnings.append('Vendor email differs from database record')
            except Exception:
                pass

        # RFQ
        rfq_no_val = detailed.get('rfq_number', ExtractedField(None)).value
        rfq_id, rfq_conf = cls.match_rfq(rfq_no_val)
        if rfq_id:
            detailed['rfq_id'] = ExtractedField(rfq_id, rfq_conf)
        elif rfq_no_val:
            warnings.append(f'RFQ "{rfq_no_val}" not found in database')

        # Match products in line items
        for item in line_items_detailed:
            pname = item.get('product_name', {}).get('value')
            sku = item.get('sku_code', {}).get('value')
            pid = cls.match_product(product_name=pname, sku=sku)
            item['matched_product_id'] = pid

        # --- Build backward-compatible flat extracted dict ---
        flat = {}
        for key, field in detailed.items():
            if isinstance(field, ExtractedField):
                flat[key] = field.value
            else:
                flat[key] = field

        # Add backward-compatible 'items' key from line items
        flat_items = []
        for li in line_items_detailed:
            flat_item = {}
            for k, v in li.items():
                if isinstance(v, dict) and 'value' in v:
                    flat_item[k] = v['value']
                else:
                    flat_item[k] = v
            # Map 'total' -> 'amount' for backward compat
            if 'total' in flat_item and 'amount' not in flat_item:
                flat_item['amount'] = flat_item['total']
            flat_items.append(flat_item)
        flat['items'] = flat_items

        # Also keep legacy keys
        if 'gst_percentage' in flat and 'tax_rate' not in flat:
            flat['tax_rate'] = flat['gst_percentage']
        if 'gst_amount' in flat and 'tax_amount' not in flat:
            flat['tax_amount'] = flat['gst_amount']

        # --- Build extracted_detailed dict ---
        ext_detailed = {}
        for key, field in detailed.items():
            if isinstance(field, ExtractedField):
                ext_detailed[key] = field.to_dict()
            else:
                ext_detailed[key] = field

        # Overall confidence
        confidence = cls._calculate_confidence(flat, line_items_detailed)

        return {
            'success': True,
            'raw_text': text,
            'confidence': confidence,
            'extracted': flat,
            'extracted_detailed': ext_detailed,
            'line_items': line_items_detailed,
            'warnings': warnings,
            'format': fmt,
            'pages': pages,
        }

    # ------------------------------------------------------------------ #
    #  Enhanced Field Extraction with Confidence                          #
    # ------------------------------------------------------------------ #

    @classmethod
    def _extract_fields_detailed(cls, text, table_data=None):
        """Extract structured fields from raw text with confidence scoring.

        Returns (dict_of_ExtractedField, list_of_warnings).
        """
        fields = {}
        warnings = []

        # --- Quote Number ---
        for pattern, conf in [
            (r'(?:quote|quotation)\s*(?:#|no\.?|number|ref\.?)[\s:]*([A-Z0-9][\w\-/]+)', 0.95),
            (r'(?:quote|quotation)\s*#?\s*:?\s*([A-Z0-9][\w\-/]+)', 0.90),
            (r'(?:proposal|estimate)\s*(?:#|no\.?|number|ref\.?)[\s:]*([A-Z0-9][\w\-/]+)', 0.85),
            (r'(?:quote|quotation)\s*#?\s*:?\s*(\d+)', 0.80),
            (r'#\s*(\d{4,})', 0.60),
        ]:
            m = re.search(pattern, text, re.IGNORECASE)
            if m:
                fields['quote_number'] = ExtractedField(
                    m.group(1).strip(), conf, m.group(0).strip()[:80]
                )
                break

        # --- Dates ---
        # Quote Date: look for date near "date" label
        date_labels = [
            (r'(?:quote|quotation)\s*date[\s:]*(.{5,30})', 0.95),
            (r'date[\s:]+(.{5,30})', 0.85),
            (r'dated[\s:]+(.{5,30})', 0.85),
        ]
        for pattern, conf in date_labels:
            m = re.search(pattern, text, re.IGNORECASE)
            if m:
                raw_date = m.group(1).strip()
                normalized = cls.normalize_date(raw_date)
                if normalized:
                    fields['quote_date'] = ExtractedField(normalized, conf, m.group(0).strip()[:80])
                    break

        # Fallback: first date found in text
        if 'quote_date' not in fields:
            for dp in cls.DATE_PATTERNS:
                dm = re.search(dp[0], text, re.IGNORECASE)
                if dm:
                    raw = dm.group(0)
                    normalized = cls.normalize_date(raw)
                    if normalized:
                        fields['quote_date'] = ExtractedField(normalized, 0.65, raw)
                        break

        # Valid Until / Validity
        validity_days = None
        for pattern, conf in [
            (r'(?:valid\s*(?:until|till|upto|up\s*to)|validity\s*date|expiry\s*date)[\s:]*(.{5,30})', 0.90),
            (r'(?:price\s*valid\s*(?:until|till))[\s:]*(.{5,30})', 0.90),
        ]:
            m = re.search(pattern, text, re.IGNORECASE)
            if m:
                raw_date = m.group(1).strip()
                normalized = cls.normalize_date(raw_date)
                if normalized:
                    fields['valid_until'] = ExtractedField(normalized, conf, m.group(0).strip()[:80])
                    break

        # Validity in days
        m = re.search(r'(?:validity|valid\s*(?:for|period))[\s:]*(\d+)\s*days?', text, re.IGNORECASE)
        if m:
            validity_days = int(m.group(1))
            if 'valid_until' not in fields and 'quote_date' in fields:
                try:
                    qd = datetime.strptime(fields['quote_date'].value, '%Y-%m-%d')
                    vu = qd + timedelta(days=validity_days)
                    fields['valid_until'] = ExtractedField(
                        vu.strftime('%Y-%m-%d'), 0.80,
                        f'Validity: {validity_days} days from quote date'
                    )
                except (ValueError, TypeError):
                    pass

        # --- Currency ---
        currency_conf = 0.0
        currency = None
        if re.search(r'\bINR\b|₹|Rs\.?\s', text):
            currency, currency_conf = 'INR', 0.92
        elif re.search(r'\bUSD\b|\$', text):
            currency, currency_conf = 'USD', 0.90
        elif re.search(r'\bEUR\b|€', text):
            currency, currency_conf = 'EUR', 0.90
        elif re.search(r'\bGBP\b|£', text):
            currency, currency_conf = 'GBP', 0.90

        if currency:
            fields['currency'] = ExtractedField(currency, currency_conf)

        # --- Amounts ---
        # Total
        for pattern, conf in [
            (r'(?:grand\s*total|net\s*payable|total\s*(?:amount|payable))[\s:]*[₹$€£]?\s*([\d,]+\.?\d*)', 0.92),
            (r'TOTAL[\s|:]*[₹$€£]?\s*([\d,]+\.?\d*)', 0.85),
            (r'(?:amount\s*due)[\s:]*[₹$€£]?\s*([\d,]+\.?\d*)', 0.80),
        ]:
            m = re.search(pattern, text, re.IGNORECASE)
            if m:
                val = m.group(1).replace(',', '')
                if cls._is_valid_amount(val):
                    fields['total_amount'] = ExtractedField(val, conf, m.group(0).strip()[:80])
                    break

        # Subtotal
        for pattern, conf in [
            (r'(?:sub\s*total|subtotal)[\s:]*[₹$€£]?\s*([\d,]+\.?\d*)', 0.88),
        ]:
            m = re.search(pattern, text, re.IGNORECASE)
            if m:
                val = m.group(1).replace(',', '')
                if cls._is_valid_amount(val):
                    fields['subtotal'] = ExtractedField(val, conf, m.group(0).strip()[:80])
                    break

        # GST / Tax
        for pattern, conf in [
            (r'GST[\s@:]*(\d+\.?\d*)\s*%', 0.92),
            (r'(\d+\.?\d*)\s*%\s*GST', 0.90),
            (r'(?:tax|vat|igst|cgst\s*\+\s*sgst)\s*(?:rate)?[\s@:]*(\d+\.?\d*)\s*%', 0.85),
            (r'(\d+\.?\d*)\s*%\s*(?:tax|vat)', 0.80),
        ]:
            m = re.search(pattern, text, re.IGNORECASE)
            if m:
                fields['gst_percentage'] = ExtractedField(m.group(1), conf, m.group(0).strip()[:80])
                break

        # GST Amount
        for pattern, conf in [
            (r'(?:GST\s*amount|tax\s*amount|GST\s*@\s*\d+%?)[\s:]*[₹$€£]?\s*([\d,]+\.?\d*)', 0.85),
            (r'(?:IGST|CGST|SGST)[\s:]*[₹$€£]?\s*([\d,]+\.?\d*)', 0.80),
        ]:
            m = re.search(pattern, text, re.IGNORECASE)
            if m:
                val = m.group(1).replace(',', '')
                if cls._is_valid_amount(val):
                    fields['gst_amount'] = ExtractedField(val, conf, m.group(0).strip()[:80])
                    break

        # --- Payment Terms ---
        for pattern, conf in [
            (r'(?:payment\s*terms?)[\s:]+(.+?)(?:\n|$)', 0.90),
            (r'((?:net|advance|immediate|cod|cash\s*on\s*delivery)\s*\d*\s*days?)', 0.85),
        ]:
            m = re.search(pattern, text, re.IGNORECASE)
            if m:
                raw_term = m.group(1).strip()[:100] if m.lastindex else m.group(0).strip()[:100]
                mapped = cls._map_payment_terms(raw_term)
                fields['payment_terms'] = ExtractedField(
                    mapped or raw_term, conf if mapped else conf - 0.15,
                    m.group(0).strip()[:80]
                )
                break

        # --- Delivery Terms ---
        for pattern, conf in [
            (r'(?:delivery\s*(?:terms?|schedule|period))[\s:]+(.+?)(?:\n|$)', 0.88),
            (r'(?:shipping\s*(?:terms?|method))[\s:]+(.+?)(?:\n|$)', 0.80),
        ]:
            m = re.search(pattern, text, re.IGNORECASE)
            if m:
                fields['delivery_terms'] = ExtractedField(
                    m.group(1).strip()[:100], conf, m.group(0).strip()[:80]
                )
                break

        # Lead Time
        for pattern, conf in [
            (r'(?:lead\s*time|delivery\s*(?:within|in))[\s:]*(\d+)\s*(?:to\s*(\d+)\s*)?(?:working\s*|business\s*)?days?', 0.88),
            (r'(\d+)\s*(?:to\s*(\d+)\s*)?(?:working\s*|business\s*)?days?\s*(?:from\s*(?:PO|order|confirmation))', 0.85),
            (r'(\d+)\s*-\s*(\d+)\s*(?:working\s*|business\s*)?days?', 0.75),
        ]:
            m = re.search(pattern, text, re.IGNORECASE)
            if m:
                # Use the larger value in a range
                val = int(m.group(2)) if m.group(2) else int(m.group(1))
                fields['lead_time_days'] = ExtractedField(val, conf, m.group(0).strip()[:80])
                break

        # Freight Terms
        for pattern, conf in [
            (r'(?:freight\s*(?:terms?|charges?))[\s:]+(.+?)(?:\n|$)', 0.88),
            (r'(?:transport(?:ation)?\s*(?:terms?|charges?))[\s:]+(.+?)(?:\n|$)', 0.80),
        ]:
            m = re.search(pattern, text, re.IGNORECASE)
            if m:
                raw_term = m.group(1).strip()[:100]
                mapped = cls._map_freight_terms(raw_term)
                fields['freight_terms'] = ExtractedField(
                    mapped or raw_term, conf if mapped else conf - 0.15,
                    m.group(0).strip()[:80]
                )
                break

        # --- Vendor Info ---
        for pattern, conf in [
            (r'(?:vendor|supplier)\s*(?:name)?[\s:]+(.+?)(?:\n|$)', 0.90),
            (r'(?:from|company)\s*(?:name)?[\s:]+(.+?)(?:\n|$)', 0.80),
            (r'(?:issued\s*by|prepared\s*by)[\s:]+(.+?)(?:\n|$)', 0.75),
        ]:
            m = re.search(pattern, text, re.IGNORECASE)
            if m:
                name = m.group(1).strip()
                # Clean up
                name = re.sub(
                    r'\s+(?:Email|Phone|Tel|Fax|Address|Contact|Date|GST|PO|Pvt|Ltd|Private|Limited).*$',
                    '', name, flags=re.IGNORECASE
                ).strip()
                if 2 <= len(name) <= 100:
                    fields['vendor_name'] = ExtractedField(name, conf, m.group(0).strip()[:80])
                    break

        # Vendor Email
        emails = re.findall(r'[\w.+-]+@[\w-]+\.[\w.]+', text)
        if emails:
            fields['vendor_email'] = ExtractedField(emails[0], 0.88)

        # Vendor Phone
        m = re.search(r'(?:phone|tel|mobile|contact)[\s:]*([+\d\s\-()]{8,20})', text, re.IGNORECASE)
        if m:
            fields['vendor_phone'] = ExtractedField(m.group(1).strip()[:20], 0.80, m.group(0).strip()[:60])

        # Vendor Address
        m = re.search(r'(?:address)[\s:]+(.+?)(?:\n\n|\n(?=[A-Z]))', text, re.IGNORECASE | re.DOTALL)
        if m:
            addr = re.sub(r'\s+', ' ', m.group(1).strip())[:200]
            if len(addr) > 5:
                fields['vendor_address'] = ExtractedField(addr, 0.75, 'Address section')

        # --- RFQ Reference ---
        for pattern, conf in [
            (r'RFQ[\s-]*(?:No|#|Ref|Number)?[\s:.]*([A-Z0-9][\w\-]+)', 0.95),
            (r'(?:reference|ref)\s*(?:no|#)?[\s:.]*\b(RFQ[\w\-]+)', 0.90),
        ]:
            m = re.search(pattern, text, re.IGNORECASE)
            if m:
                fields['rfq_number'] = ExtractedField(m.group(1).strip(), conf, m.group(0).strip()[:80])
                break

        # --- Notes / Remarks ---
        m = re.search(r'(?:notes?|remarks?|comments?)[\s:]+(.+?)(?:\n\n|$)', text, re.IGNORECASE | re.DOTALL)
        if m:
            notes = m.group(1).strip()[:300]
            if len(notes) > 3:
                fields['notes'] = ExtractedField(notes, 0.70)

        return fields, warnings

    # ------------------------------------------------------------------ #
    #  Line Item Extraction                                               #
    # ------------------------------------------------------------------ #

    @classmethod
    def _extract_line_items(cls, text, table_data=None):
        """Extract line items from text and/or structured table data.

        Returns list of dicts, each with per-field {value, confidence} entries.
        """
        items = []

        # Strategy 1: Extract from pdfplumber table data (highest confidence)
        if table_data:
            table_items = cls._extract_from_tables(table_data)
            if table_items:
                items.extend(table_items)

        # Strategy 2: Text-based table parsing (detect header row, map columns)
        if not items:
            table_text_items = cls._extract_from_text_table(text)
            items.extend(table_text_items)

        # Strategy 3: Parse from text lines (column alignment / pipe-separated)
        if not items:
            text_items = cls._extract_from_text_lines(text)
            items.extend(text_items)

        # Strategy 4: Multi-line chat/WhatsApp format (product code + description + price on separate lines)
        if not items:
            chat_items = cls._extract_from_chat_multiline(text)
            items.extend(chat_items)

        return items

    @classmethod
    def _extract_from_text_table(cls, text):
        """Detect table headers in text and extract rows based on column positions.

        Handles formats like:
        Product Name  Specification  Quantity  Unit  Rate  Tax Percentage  Taxable Amount
        Carton Box    (EXPORT)       400.000   Nos   63.50  GST5          ₹ 25,400.00
        """
        items = []
        lines = text.split('\n')

        # Column keyword mapping
        col_keywords = {
            'product': ['product name', 'product', 'description', 'item', 'material', 'particulars', 'description of goods'],
            'spec': ['specification', 'spec', 'details'],
            'qty': ['quantity', 'qty', 'nos', 'no.of'],
            'uom': ['unit', 'uom', 'units'],
            'price': ['rate', 'price', 'unit price', 'unit rate', 'cost', 'per'],
            'tax': ['tax percentage', 'tax%', 'gst', 'gst%', 'tax rate', 'igst'],
            'hsn': ['hsn', 'hsn/sac', 'hsn code', 'sac'],
            'total': ['taxable amount', 'amount', 'total', 'value', 'net amount', 'line total'],
            'discount': ['discount', 'disc', 'disc%', 'disc.'],
        }

        header_idx = -1
        col_positions = {}  # {field: (start_pos, end_pos)}

        # Find the header row
        for idx, line in enumerate(lines):
            line_lower = line.lower().strip()
            if not line_lower:
                continue

            matches = 0
            temp_positions = {}
            for field, keywords in col_keywords.items():
                for kw in keywords:
                    pos = line_lower.find(kw)
                    if pos >= 0:
                        if field not in temp_positions:
                            temp_positions[field] = pos
                            matches += 1
                        break

            if matches >= 3:  # Need at least 3 column matches
                header_idx = idx
                # Sort by position to determine column boundaries
                sorted_cols = sorted(temp_positions.items(), key=lambda x: x[1])
                for i, (field, pos) in enumerate(sorted_cols):
                    end = sorted_cols[i + 1][1] if i + 1 < len(sorted_cols) else len(line) + 50
                    col_positions[field] = (pos, end)
                break

        if header_idx < 0 or 'product' not in col_positions:
            return items

        # Extract data rows after header
        data_lines = lines[header_idx + 1:]
        current_product = None
        current_item = None

        for line in data_lines:
            if not line.strip():
                if current_item:
                    # Empty line may end the current product
                    continue
                continue

            line_lower = line.lower().strip()

            # Skip total/summary rows
            if re.match(r'^\s*(total|sub\s*total|grand|net|amount|terms|note|condition|disclaimer)', line_lower):
                break

            # Check if this line has numeric data (qty, price, amount)
            has_numbers = bool(re.search(r'\d+\.?\d*', line))
            prod_start, prod_end = col_positions.get('product', (0, 30))

            # Extract text at each column position
            row_data = {}
            for field, (start, end) in col_positions.items():
                cell = line[start:end].strip() if start < len(line) else ''
                cell = cell.strip('₹ \t|').strip()
                if cell:
                    row_data[field] = cell

            # If we have product text AND numbers, it's a data row
            product_text = row_data.get('product', '').strip()
            qty_text = row_data.get('qty', '')
            price_text = row_data.get('price', '')
            total_text = row_data.get('total', '')

            qty_num = cls._parse_number(qty_text)
            price_num = cls._parse_number(price_text)
            total_num = cls._parse_number(total_text)

            # FALLBACK: If column-position mapping doesn't yield numbers,
            # try extracting "QTY UOM PRICE TOTAL" from the right side of the line
            if not (qty_num is not None and price_num is not None) and has_numbers:
                right_match = re.search(
                    r'(\d+\.?\d{0,3})\s+(Nos|Pcs|KG|MTS|LTR|BOX|NOS|PCS|MT|Sets?|Pairs?|Units?)\s+'
                    r'([\d,.]+)\s+(?:GST\d*\s+)?(?:₹\s*)?([\d,.]+)',
                    line, re.IGNORECASE
                )
                if right_match:
                    qty_num = cls._parse_number(right_match.group(1))
                    row_data['uom'] = right_match.group(2).upper().rstrip('.')
                    price_num = cls._parse_number(right_match.group(3))
                    total_num = cls._parse_number(right_match.group(4))
                    # Product name = everything BEFORE the matched numbers
                    prod_text = line[:right_match.start()].strip()
                    if prod_text:
                        product_text = prod_text

            if product_text and (qty_num is not None or price_num is not None or total_num is not None):
                # New product row
                if current_item and current_item['product_name']['value']:
                    items.append(current_item)

                current_item = cls._make_empty_line_item()
                current_item['product_name'] = {'value': product_text, 'confidence': 0.88}
                if qty_num is not None:
                    current_item['quantity'] = {'value': qty_num, 'confidence': 0.90}
                if price_num is not None:
                    current_item['unit_price'] = {'value': price_num, 'confidence': 0.90}
                if total_num is not None:
                    current_item['total'] = {'value': total_num, 'confidence': 0.88}
                uom = row_data.get('uom', '').strip().upper().rstrip('.')
                if uom:
                    current_item['uom'] = {'value': uom, 'confidence': 0.85}
                spec = row_data.get('spec', '')
                if spec:
                    current_item['description'] = {'value': spec, 'confidence': 0.80}
                hsn = row_data.get('hsn', '')
                if hsn:
                    current_item['sku_code'] = {'value': hsn, 'confidence': 0.75}
                tax = row_data.get('tax', '')
                tax_num = cls._parse_number(re.sub(r'[^0-9.]', '', tax)) if tax else None
                if tax_num is not None:
                    current_item['gst_rate'] = {'value': tax_num, 'confidence': 0.82}
                disc = row_data.get('discount', '')
                disc_num = cls._parse_number(disc)
                if disc_num is not None:
                    current_item['discount'] = {'value': disc_num, 'confidence': 0.80}

            elif product_text and current_item and not qty_num and not price_num:
                # Continuation line (multi-line product name/spec)
                existing = current_item['product_name']['value'] or ''
                current_item['product_name']['value'] = (existing + ' ' + product_text).strip()
                spec = row_data.get('spec', '')
                if spec:
                    existing_spec = current_item['description']['value'] or ''
                    current_item['description']['value'] = (existing_spec + ' ' + spec).strip()

        # Don't forget the last item
        if current_item and current_item['product_name']['value']:
            items.append(current_item)

        return items

    @classmethod
    def _extract_from_tables(cls, table_data):
        """Extract line items from structured table data (e.g. pdfplumber tables)."""
        items = []

        for table in table_data:
            if not table or len(table) < 2:
                continue

            # Find header row — look for rows containing keywords
            header_idx = -1
            col_map = {}
            header_keywords = {
                'product': ['product', 'item', 'description', 'material', 'name', 'particulars'],
                'sku': ['sku', 'code', 'part no', 'part#', 'item code', 'material code'],
                'qty': ['qty', 'quantity', 'nos', 'no.'],
                'uom': ['uom', 'unit', 'units'],
                'price': ['price', 'rate', 'unit price', 'unit rate', 'cost'],
                'discount': ['discount', 'disc', 'disc%'],
                'gst': ['gst', 'tax', 'gst%', 'tax%', 'igst', 'gst rate'],
                'total': ['total', 'amount', 'value', 'line total', 'net amount'],
            }

            for idx, row in enumerate(table):
                if not row:
                    continue
                row_str = ' '.join(str(c or '').lower() for c in row)
                matches = 0
                temp_map = {}
                for col_idx, cell in enumerate(row):
                    cell_lower = str(cell or '').lower().strip()
                    for field, keywords in header_keywords.items():
                        if any(kw in cell_lower for kw in keywords):
                            if field not in temp_map:
                                temp_map[field] = col_idx
                                matches += 1
                if matches >= 2:
                    header_idx = idx
                    col_map = temp_map
                    break

            if header_idx < 0:
                continue

            # Parse data rows
            for row in table[header_idx + 1:]:
                if not row:
                    continue
                # Skip rows that are all empty/None
                if all(not str(c or '').strip() for c in row):
                    continue

                # Skip summary rows
                row_text = ' '.join(str(c or '').lower() for c in row)
                if re.search(r'(?:sub\s*total|grand\s*total|total|tax|gst\s*amount)', row_text):
                    continue

                item = cls._make_empty_line_item()

                for field, col_idx in col_map.items():
                    if col_idx < len(row):
                        raw_val = str(row[col_idx] or '').strip()
                        if not raw_val:
                            continue

                        if field == 'product':
                            item['product_name'] = {'value': raw_val, 'confidence': 0.92}
                        elif field == 'sku':
                            item['sku_code'] = {'value': raw_val, 'confidence': 0.95}
                        elif field == 'qty':
                            num = cls._parse_number(raw_val)
                            if num is not None:
                                item['quantity'] = {'value': num, 'confidence': 0.93}
                        elif field == 'uom':
                            item['uom'] = {'value': raw_val.upper(), 'confidence': 0.90}
                        elif field == 'price':
                            num = cls._parse_number(raw_val)
                            if num is not None:
                                item['unit_price'] = {'value': num, 'confidence': 0.92}
                        elif field == 'discount':
                            num = cls._parse_number(raw_val)
                            if num is not None:
                                item['discount'] = {'value': num, 'confidence': 0.85}
                        elif field == 'gst':
                            num = cls._parse_number(raw_val)
                            if num is not None:
                                item['gst_rate'] = {'value': num, 'confidence': 0.88}
                        elif field == 'total':
                            num = cls._parse_number(raw_val)
                            if num is not None:
                                item['total'] = {'value': num, 'confidence': 0.92}

                # Only keep if we got at least a product name or qty+price
                if (item['product_name']['value'] or
                        (item['quantity']['value'] is not None and item['unit_price']['value'] is not None)):
                    items.append(item)

        return items

    @classmethod
    def _extract_from_text_lines(cls, text):
        """Extract line items from text using pattern matching on each line."""
        items = []
        lines = text.split('\n')

        skip_keywords = re.compile(
            r'^(?:#|sr|s\.?no|sl|description|item|product|service|qty|quantity|'
            r'price|amount|total|subtotal|sub\s*total|tax|gst|grand|net|'
            r'quotation|quote|rfq|date|vendor|email|payment|delivery|validity|'
            r'terms|unit\s*price|uom|sku|response|header|particulars)',
            re.IGNORECASE,
        )

        for line in lines:
            line = line.strip()
            if not line or len(line) < 5:
                continue
            if skip_keywords.match(line):
                continue

            item = cls._try_parse_line_item(line)
            if item:
                items.append(item)

        return items

    @classmethod
    def _try_parse_line_item(cls, line):
        """Try various patterns to parse a single line as a line item.

        Returns a dict with per-field {value, confidence} or None.
        """
        # Pattern 0: Invoice with PARTNO/HSN — "# PARTNO/ HSN DESCRIPTION QTY TAX% RATE AMOUNT"
        # e.g. "1 38459582/ 34039900 ULTRA COOLANT 20 LTS 2 18% 51837.0 103674.00"
        # Strategy: extract the LAST numbers (amount, rate, tax%, qty) and treat everything before as description
        m = re.match(
            r'(\d+)\s+'                              # Sl. No
            r'([\d\w]{5,15}/?)\s+'                   # Part No
            r'(\d{6,8})\s+'                          # HSN Code
            r'(.+?)\s+'                              # Description (non-greedy)
            r'(\d{1,6})\s+'                          # Quantity (1-6 digits, must be small number)
            r'(\d{1,3})\s*%\s+'                      # Tax Rate % (1-3 digits + %)
            r'([\d,.]+)\s+'                          # Unit Rate
            r'([\d,.]+)',                             # Amount
            line,
        )
        if m:
            desc = m.group(4).strip()
            qty = cls._parse_number(m.group(5))
            gst = cls._parse_number(m.group(6))
            rate = cls._parse_number(m.group(7))
            amount = cls._parse_number(m.group(8))
            if rate and amount and (qty or 0) < 100000:
                item = cls._make_empty_line_item()
                item['product_name'] = {'value': desc, 'confidence': 0.92}
                item['sku_code'] = {'value': m.group(2).strip().rstrip('/'), 'confidence': 0.92}
                item['quantity'] = {'value': qty, 'confidence': 0.92}
                item['unit_price'] = {'value': rate, 'confidence': 0.92}
                item['total'] = {'value': amount, 'confidence': 0.92}
                if gst:
                    item['gst_rate'] = {'value': gst, 'confidence': 0.90}
                return item

        # Pattern 0 fallback: Try extracting from the END of the line
        # Look for RATE and AMOUNT as the last two numbers, TAX% before them
        # Works for: "# PARTNO/ HSN DESCRIPTION QTY TAX% RATE AMOUNT"
        m = re.match(r'(\d+)\s+([\d\w]{5,15}/?)\s+(\d{6,8})\s+(.+)', line)
        if m:
            rest = m.group(4).strip()
            # Extract last 2-3 numbers from the end
            nums = list(re.finditer(r'([\d,.]+\.?\d*)', rest))
            if len(nums) >= 2:
                amount_match = nums[-1]
                rate_match = nums[-2]
                amount = cls._parse_number(amount_match.group(1))
                rate = cls._parse_number(rate_match.group(1))

                # Find tax% (number followed by %)
                tax_m = re.search(r'(\d{1,3})\s*%', rest)
                gst = cls._parse_number(tax_m.group(1)) if tax_m else None

                # Find qty (small number before tax%)
                qty = None
                if tax_m:
                    before_tax = rest[:tax_m.start()].strip()
                    qty_m = re.search(r'(\d{1,5})\s*$', before_tax)
                    if qty_m:
                        qty = cls._parse_number(qty_m.group(1))
                        desc = before_tax[:qty_m.start()].strip()
                    else:
                        desc = before_tax
                elif len(nums) >= 3:
                    qty_match = nums[-3]
                    qty = cls._parse_number(qty_match.group(1))
                    desc = rest[:qty_match.start()].strip()
                else:
                    desc = rest[:rate_match.start()].strip()

                if rate and amount and rate < 10000000 and amount < 100000000 and desc:
                    item = cls._make_empty_line_item()
                    item['product_name'] = {'value': desc, 'confidence': 0.85}
                    item['sku_code'] = {'value': m.group(2).strip().rstrip('/'), 'confidence': 0.88}
                    if qty and qty < 100000:
                        item['quantity'] = {'value': qty, 'confidence': 0.85}
                    item['unit_price'] = {'value': rate, 'confidence': 0.85}
                    item['total'] = {'value': amount, 'confidence': 0.85}
                    if gst:
                        item['gst_rate'] = {'value': gst, 'confidence': 0.82}
                    return item

        # Pattern 1: # | Product | SKU | Qty | UOM | UnitPrice | Total
        m = re.match(
            r'(\d+)\s*[|\s]+(.+?)\s+([A-Z]{2,5}[\-]?\d{2,}[\w\-]*)\s+(\d+\.?\d*)\s+([A-Z]{1,5})\s+(\d+[\d,.]*)\s+(\d+[\d,.]*)',
            line,
        )
        if m:
            item = cls._make_empty_line_item()
            item['product_name'] = {'value': m.group(2).strip(), 'confidence': 0.88}
            item['sku_code'] = {'value': m.group(3).strip(), 'confidence': 0.92}
            item['quantity'] = {'value': cls._parse_number(m.group(4)), 'confidence': 0.90}
            item['uom'] = {'value': m.group(5).strip(), 'confidence': 0.88}
            item['unit_price'] = {'value': cls._parse_number(m.group(6)), 'confidence': 0.88}
            item['total'] = {'value': cls._parse_number(m.group(7)), 'confidence': 0.88}
            return item

        # Pattern 2: # | Product | Qty | UnitPrice | Total
        m = re.match(
            r'(\d+)\s*[|\s]+(.{3,50}?)\s*[|\s]+(\d+\.?\d*)\s*[|\s]+[$₹]?\s*(\d+[\d,.]*)\s*[|\s]+[$₹]?\s*(\d+[\d,.]*)',
            line,
        )
        if m:
            desc = m.group(2).strip()
            qty = cls._parse_number(m.group(3))
            price = cls._parse_number(m.group(4))
            amount = cls._parse_number(m.group(5))
            if amount and amount > 0 and price and price > 0:
                item = cls._make_empty_line_item()
                item['product_name'] = {'value': desc, 'confidence': 0.82}
                item['quantity'] = {'value': qty, 'confidence': 0.85}
                item['unit_price'] = {'value': price, 'confidence': 0.85}
                item['total'] = {'value': amount, 'confidence': 0.85}
                return item

        # Pattern 3: Pipe-separated generic
        parts = [p.strip() for p in line.split('|') if p.strip()]
        if len(parts) >= 4:
            try:
                nums = []
                for p in reversed(parts):
                    cleaned = p.replace(',', '').replace('$', '').replace('₹', '').strip()
                    float(cleaned)
                    nums.insert(0, cleaned)
                    if len(nums) >= 3:
                        break
                if len(nums) >= 2:
                    desc_parts = parts[:len(parts) - len(nums)]
                    desc = ' '.join(desc_parts).strip()
                    if desc and len(desc) > 2:
                        item = cls._make_empty_line_item()
                        item['product_name'] = {'value': desc, 'confidence': 0.75}
                        if len(nums) == 3:
                            item['quantity'] = {'value': cls._parse_number(nums[0]), 'confidence': 0.78}
                            item['unit_price'] = {'value': cls._parse_number(nums[1]), 'confidence': 0.78}
                            item['total'] = {'value': cls._parse_number(nums[2]), 'confidence': 0.78}
                        elif len(nums) == 2:
                            item['unit_price'] = {'value': cls._parse_number(nums[0]), 'confidence': 0.72}
                            item['total'] = {'value': cls._parse_number(nums[1]), 'confidence': 0.72}
                        return item
            except ValueError:
                pass

        # Pattern 4: Space-separated "Description  Price  Qty  Amount"
        m = re.match(
            r'(.{5,40}?)\s{2,}(\d+[\d,.]*)\s{2,}(\d+\.?\d*)\s{2,}(\d+[\d,.]*)',
            line,
        )
        if m:
            desc = m.group(1).strip()
            if len(desc) > 2:
                price = cls._parse_number(m.group(2))
                qty = cls._parse_number(m.group(3))
                amount = cls._parse_number(m.group(4))
                if amount and amount > 0:
                    item = cls._make_empty_line_item()
                    item['product_name'] = {'value': desc, 'confidence': 0.70}
                    item['quantity'] = {'value': qty, 'confidence': 0.70}
                    item['unit_price'] = {'value': price, 'confidence': 0.70}
                    item['total'] = {'value': amount, 'confidence': 0.70}
                    return item

        # Pattern 5: "Model: CODE (description) - QTY no - Rs.PRICE/pc"
        # e.g. "Model: R960-220000-000 (steering axle assembly) - 1 no - Rs.77180/pc"
        m = re.match(
            r'(?:Model|Part|Item)[\s:]+([A-Z0-9][\w\-./]+)\s*[\(\[]?([^)\]]*?)[\)\]]?\s*[-–]\s*(\d+\.?\d*)\s*(?:nos?|pcs?|pc|units?|set)?\s*[-–]\s*(?:Rs\.?|INR|₹)\s*([\d,.]+)',
            line, re.IGNORECASE,
        )
        if m:
            sku = m.group(1).strip()
            desc = m.group(2).strip() or sku
            item = cls._make_empty_line_item()
            item['product_name'] = {'value': desc, 'confidence': 0.88}
            item['sku_code'] = {'value': sku, 'confidence': 0.92}
            item['quantity'] = {'value': cls._parse_number(m.group(3)), 'confidence': 0.90}
            item['unit_price'] = {'value': cls._parse_number(m.group(4)), 'confidence': 0.90}
            item['uom'] = {'value': 'NOS', 'confidence': 0.70}
            return item

        # Pattern 6: "QTY UNIT PRODUCT Rs.PRICE" or "QTY PRODUCT PRICE each"
        # e.g. "50 bags rice 2000 each", "25 kg sugar Rs 1200", "Need 50 bags rice Rs 2000 each"
        m = re.match(
            r'(?:Need\s+|Also\s+|And\s+|Plus\s+)?(\d+\.?\d*)\s*(bags?|pcs?|nos?|kg|kgs?|mts?|ltrs?|ltr|boxes?|drums?|sets?|pairs?|rolls?|bundles?|units?)?\s+(.+?)\s+(?:Rs\.?|INR|₹|@)?\s*([\d,.]+)\s*(?:/?\s*(?:pc|each|per|unit|kg|nos?)?)?(?:\s*=\s*(?:Rs\.?|INR|₹)?\s*([\d,.]+))?',
            line, re.IGNORECASE,
        )
        if m:
            qty = cls._parse_number(m.group(1))
            uom = (m.group(2) or '').upper().rstrip('S') or 'NOS'
            uom_map = {'BAG': 'BAG', 'PC': 'PCS', 'NO': 'NOS', 'KG': 'KG', 'MT': 'MTS', 'LTR': 'LTR', 'BOX': 'BOX', 'DRUM': 'DRUM', 'SET': 'SET', 'PAIR': 'PAIR', 'ROLL': 'ROLL', 'BUNDLE': 'BUNDLE', 'UNIT': 'NOS'}
            uom = uom_map.get(uom, uom)
            desc = m.group(3).strip().rstrip('-–')
            price = cls._parse_number(m.group(4))
            total = cls._parse_number(m.group(5)) if m.group(5) else (qty * price if qty and price else None)
            if desc and len(desc) > 1:
                item = cls._make_empty_line_item()
                item['product_name'] = {'value': desc, 'confidence': 0.82}
                item['quantity'] = {'value': qty, 'confidence': 0.85}
                item['uom'] = {'value': uom, 'confidence': 0.80}
                item['unit_price'] = {'value': price, 'confidence': 0.85}
                if total:
                    item['total'] = {'value': total, 'confidence': 0.80}
                return item

        # Pattern 7: "PRODUCT QTY.00 UOM PRICE.00 UOM DISC% AMOUNT"
        # e.g. "Reed Switch AM40-0-FL-04 ... 2.00 Nos. 1,027.00 Nos. 18% 1,684.28"
        m = re.match(
            r'(?:\d+\s+)?(.+?)\s+(\d+\.?\d*)\s*(?:Nos\.?|Pcs\.?|KG|MTS|LTR|NOS|PCS|BOX)\s+([\d,.]+)\s*(?:Nos\.?|Pcs\.?|KG|MTS|LTR|NOS|PCS|BOX)?\s*(?:(\d+\.?\d*)\s*%\s*)?([\d,.]+)',
            line, re.IGNORECASE,
        )
        if m:
            desc = m.group(1).strip()
            # Filter out header rows
            if not re.match(r'^(?:Sl|No|#|Sr|S\.)', desc, re.IGNORECASE) and len(desc) > 2:
                qty = cls._parse_number(m.group(2))
                price = cls._parse_number(m.group(3))
                disc = cls._parse_number(m.group(4)) if m.group(4) else None
                total = cls._parse_number(m.group(5))
                if price and price > 0 and total and total > 0:
                    item = cls._make_empty_line_item()
                    item['product_name'] = {'value': desc, 'confidence': 0.82}
                    item['quantity'] = {'value': qty, 'confidence': 0.85}
                    item['unit_price'] = {'value': price, 'confidence': 0.85}
                    item['total'] = {'value': total, 'confidence': 0.85}
                    if disc:
                        item['discount'] = {'value': disc, 'confidence': 0.80}
                    # Try to extract UOM
                    uom_m = re.search(r'(\d+\.?\d*)\s*(Nos?|Pcs?|KG|MTS|LTR|BOX)', line, re.IGNORECASE)
                    if uom_m:
                        item['uom'] = {'value': uom_m.group(2).upper().rstrip('.'), 'confidence': 0.80}
                    return item

        # Pattern 8: "ITEM  QTY  TOTAL" (receipt/transport format)
        # e.g. "1 C BOX MED   1   120.00"
        m = re.match(
            r'(?:\d+\s+)?(.{3,40}?)\s{2,}(\d+\.?\d*)\s{2,}([\d,.]+)',
            line,
        )
        if m:
            desc = m.group(1).strip()
            # Filter out total/summary rows and headers
            if not re.match(r'(?:total|sub|grand|freight|loading|unloading|sgst|cgst|igst|gst|tax|disc|round|net|amount)', desc, re.IGNORECASE):
                qty = cls._parse_number(m.group(2))
                total = cls._parse_number(m.group(3))
                if qty and total and total > 0 and len(desc) > 2:
                    item = cls._make_empty_line_item()
                    item['product_name'] = {'value': desc, 'confidence': 0.68}
                    item['quantity'] = {'value': qty, 'confidence': 0.72}
                    item['total'] = {'value': total, 'confidence': 0.72}
                    if qty > 0:
                        item['unit_price'] = {'value': round(total / qty, 2), 'confidence': 0.60}
                    return item

        return None

    @classmethod
    def _extract_from_chat_multiline(cls, text):
        """Extract products from multi-line WhatsApp/chat format.

        Handles patterns like:
        AM41_1_FL_04
        Turnable magnetic sensor with flying lead reed switch 3 wire
        Rs.1027/-
        Discount 15%
        GST 18%
        Delivery 3 weeks
        """
        items = []
        lines = text.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()

            # Look for a product code pattern (alphanumeric with underscores/dashes)
            code_match = re.match(r'^([A-Z0-9][A-Z0-9_\-./]{3,30})$', line, re.IGNORECASE)
            if code_match:
                sku = code_match.group(1)
                desc = ''
                price = None
                discount = None
                gst = None
                qty = 1

                # Scan next few lines for description, price, discount, GST
                for j in range(1, min(8, len(lines) - i)):
                    next_line = lines[i + j].strip()
                    if not next_line:
                        continue

                    # Price: Rs.1027/- or Rs 1027 or ₹1027
                    pm = re.match(r'(?:Rs\.?|INR|₹)\s*([\d,.]+)(?:/[-–])?', next_line, re.IGNORECASE)
                    if pm:
                        price = cls._parse_number(pm.group(1))
                        continue

                    # Discount
                    dm = re.match(r'Discount\s*([\d.]+)\s*%', next_line, re.IGNORECASE)
                    if dm:
                        discount = cls._parse_number(dm.group(1))
                        continue

                    # GST
                    gm = re.match(r'GST\s*([\d.]+)\s*%', next_line, re.IGNORECASE)
                    if gm:
                        gst = cls._parse_number(gm.group(1))
                        continue

                    # Quantity: "2 NOS" or "Qty: 5"
                    qm = re.match(r'(?:Qty|Quantity)?[\s:]*(\d+)\s*(?:NOS|PCS|KG|NOS?|nos)', next_line, re.IGNORECASE)
                    if qm:
                        qty = cls._parse_number(qm.group(1))
                        continue

                    # Delivery/notes — skip
                    if re.match(r'(?:Delivery|Note|Regard|Dear|OK)', next_line, re.IGNORECASE):
                        continue

                    # If it looks like another code, stop
                    if re.match(r'^[A-Z0-9][A-Z0-9_\-./]{3,30}$', next_line, re.IGNORECASE):
                        break

                    # Otherwise it's likely the description
                    if not desc and len(next_line) > 5 and not re.match(r'^\d', next_line):
                        desc = next_line

                if price:
                    item = cls._make_empty_line_item()
                    item['sku_code'] = {'value': sku, 'confidence': 0.90}
                    item['product_name'] = {'value': desc or sku, 'confidence': 0.85 if desc else 0.70}
                    item['unit_price'] = {'value': price, 'confidence': 0.88}
                    item['quantity'] = {'value': qty, 'confidence': 0.75}
                    item['uom'] = {'value': 'NOS', 'confidence': 0.65}
                    if discount:
                        item['discount'] = {'value': discount, 'confidence': 0.85}
                    if gst:
                        item['gst_rate'] = {'value': gst, 'confidence': 0.85}
                    items.append(item)

            i += 1
        return items

    # ------------------------------------------------------------------ #
    #  Helpers                                                            #
    # ------------------------------------------------------------------ #

    @classmethod
    def _make_empty_line_item(cls):
        """Create empty line item template with all fields."""
        return {
            'product_name': {'value': '', 'confidence': 0.0},
            'sku_code': {'value': '', 'confidence': 0.0},
            'quantity': {'value': None, 'confidence': 0.0},
            'uom': {'value': '', 'confidence': 0.0},
            'unit_price': {'value': None, 'confidence': 0.0},
            'discount': {'value': None, 'confidence': 0.0},
            'gst_rate': {'value': None, 'confidence': 0.0},
            'total': {'value': None, 'confidence': 0.0},
            'description': {'value': '', 'confidence': 0.0},
            'matched_product_id': None,
        }

    @classmethod
    def _parse_number(cls, s):
        """Parse a number string, removing commas and currency symbols."""
        if s is None:
            return None
        s = str(s).replace(',', '').replace('₹', '').replace('$', '').replace('€', '').replace('£', '').strip()
        try:
            val = float(s)
            # Return int if it's a whole number
            if val == int(val) and '.' not in str(s):
                return int(val)
            return val
        except (ValueError, TypeError):
            return None

    @classmethod
    def _is_valid_amount(cls, s):
        """Check if string is a valid monetary amount."""
        try:
            v = float(str(s).replace(',', ''))
            return v > 0
        except (ValueError, TypeError):
            return False

    @classmethod
    def _map_payment_terms(cls, raw):
        """Map raw payment terms text to standard code."""
        if not raw:
            return None
        lower = raw.lower().strip()
        for key, code in cls.PAYMENT_TERMS_MAP.items():
            if key in lower:
                return code
        # Try extracting "net N" pattern
        m = re.search(r'net\s*(\d+)', lower)
        if m:
            days = m.group(1)
            code = f'NET_{days}'
            if code in ('NET_7', 'NET_15', 'NET_30', 'NET_45', 'NET_60', 'NET_90'):
                return code
        return None

    @classmethod
    def _map_freight_terms(cls, raw):
        """Map raw freight terms text to standard code."""
        if not raw:
            return None
        lower = raw.lower().strip()
        for key, code in cls.FREIGHT_TERMS_MAP.items():
            if key in lower:
                return code
        return None

    # ------------------------------------------------------------------ #
    #  Confidence Calculation                                             #
    # ------------------------------------------------------------------ #

    @classmethod
    def _calculate_confidence(cls, flat_extracted, line_items):
        """Calculate overall extraction confidence (0-100).

        Weights key fields that indicate a valid quote was parsed.
        """
        score = 0
        total_weight = 8.0

        if flat_extracted.get('quote_number'):
            score += 1.0
        if flat_extracted.get('quote_date'):
            score += 1.0
        if flat_extracted.get('total_amount'):
            score += 1.5
        if flat_extracted.get('vendor_name'):
            score += 1.0
        if flat_extracted.get('currency'):
            score += 0.5
        if line_items and len(line_items) > 0:
            score += 2.0
        if flat_extracted.get('payment_terms') or flat_extracted.get('delivery_terms'):
            score += 0.5
        if flat_extracted.get('gst_percentage') or flat_extracted.get('gst_amount'):
            score += 0.5

        return min(round(score / total_weight * 100), 100)
